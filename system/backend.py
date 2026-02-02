import os
import re
import random
import httpx


import sys
import ctypes
import json
import shutil
import sqlite3
import logging
import time
import subprocess
import uuid
from pathlib import Path
from typing import List, Optional, Dict, Any, Generator
from datetime import datetime, timedelta
from contextlib import asynccontextmanager
import gc

try:
    import psutil
except ImportError:
    psutil = None

from fastapi import FastAPI, HTTPException, Depends, BackgroundTasks, status, Body, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.responses import StreamingResponse, FileResponse
from pydantic import BaseModel
from passlib.context import CryptContext
from jose import JWTError, jwt
import asyncio

# Discord Bot Integration
# Discord Bot Removed


# Llama.cpp
try:
    from llama_cpp import Llama
except ImportError:
    Llama = None

# ChromaDB
try:
    import chromadb
    from chromadb.config import Settings
except ImportError:
    chromadb = None

# Document Loaders
try:
    import docx
except ImportError:
    docx = None
try:
    import openpyxl
except ImportError:
    openpyxl = None

try:
    import pypdf
except ImportError:
    pypdf = None

import hashlib

# Setup Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("oonanji-backend")

# --- Configuration ---
BASE_DIR = Path(__file__).parent
MODELS_DIR = BASE_DIR / "models"
MNT_DIR = BASE_DIR / "mnt"
INTERNAL_NAS_DIR = BASE_DIR / "internal_storage"
CHROMA_DB_DIR = BASE_DIR / "chroma_db"
DB_PATH = BASE_DIR / "users.db"

# Safety check: if users.db is a directory (sometimes happens with Docker volume mounts), remove it
if DB_PATH.exists() and DB_PATH.is_dir():
    logger.warning("DB_PATH is a directory, removing it to allow file creation...")
    shutil.rmtree(DB_PATH)

# Ensure directories exist
MODELS_DIR.mkdir(exist_ok=True)
# MNT_DIR is for external NAS mounts, do not auto-create to respect user's filesystem
INTERNAL_NAS_DIR.mkdir(exist_ok=True)

# Security
import secrets
KEY_FILE = BASE_DIR / "secret.key"
if KEY_FILE.exists():
    SECRET_KEY = KEY_FILE.read_text().strip()
else:
    SECRET_KEY = secrets.token_urlsafe(32)
    KEY_FILE.write_text(SECRET_KEY)

ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 30  # 30 days
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# Global State
class GlobalState:
    is_indexing = False
    indexing_progress = 0.0
    indexing_status = "Idle"
    indexing_log: List[str] = []
    stop_indexing_flag: bool = False
    current_storage_mode = "nas"  # 'nas' or 'internal'
    indexing_start_time = None
    indexing_total_files = 0
    indexing_processed_files = 0
    last_indexed_at: Optional[str] = None

state = GlobalState()

def get_db_status():
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'indexing_status'")
        result = cursor.fetchone()
        conn.close()
        if result:
            return json.loads(result[0])
        return {}
    except Exception:
        return {}

def get_storage_mode():
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'storage_mode'")
        result = cursor.fetchone()
        conn.close()
        if result:
            return result[0]
        return "nas" # Default
    except Exception:
        return "nas"

def ensure_user_models_dir(username: str) -> Path:
    if username == "adminuser":
        return MODELS_DIR

    user_model_dir = BASE_DIR / f"models_{username}"
    
    # Check if needs migration from symlink to directory
    if user_model_dir.is_symlink() or (user_model_dir.exists() and not user_model_dir.is_dir()):
        logger.info(f"Removing existing symlink/file for {username} models to replace with copy...")
        try:
            os.unlink(user_model_dir)
        except Exception as e:
            logger.error(f"Failed to remove symlink for {username}: {e}")

    if not user_model_dir.exists():
        logger.info(f"Creating model directory for {username}...")
        try:
            if not MODELS_DIR.exists():
                logger.warning("Base models directory missing, creating empty user dir")
                user_model_dir.mkdir(parents=True, exist_ok=True)
            else:
                # USER REQUEST: Physical copy of models for each user
                # Warning: This consumes significant disk space
                logger.info(f"Copying models for {username} (Physical Copy)...")
                shutil.copytree(MODELS_DIR, user_model_dir, dirs_exist_ok=True)
                
        except Exception as e:
            logger.error(f"Failed to setup models for {username}: {e}")
            return MODELS_DIR
    return user_model_dir



class RemoteLlama:
    def __init__(self, base_url: str, model_path: str):
        self.base_url = base_url.rstrip('/')
        self.model_path = model_path
        
    def create_chat_completion(self, messages, max_tokens=1024, temperature=0.7, stream=True, **kwargs):
        url = f"{self.base_url}/v1/chat/completions"
        headers = {"Content-Type": "application/json"}
        
        # Extract system prompt if present to map to "model" if needed, 
        # but standard OpenAI API just takes messages. 
        # We pass a generic model name or specific one if the worker supports it.
        payload = {
            "model": "default", # Worker should handle the model or we pass the filename if mapped
            "messages": messages,
            "max_tokens": max_tokens,
            "temperature": temperature,
            "stream": stream
        }
        
        if stream:
            with httpx.stream("POST", url, json=payload, headers=headers, timeout=60.0) as response:
                if response.status_code != 200:
                    logger.error(f"Remote Worker Error: {response.text}")
                    yield {"choices": [{"delta": {"content": f" Error: Remote worker returned {response.status_code}"}}]}
                    return

                for line in response.iter_lines():
                    if line.startswith("data: "):
                        data_str = line[6:]
                        if data_str.strip() == "[DONE]":
                            break
                        try:
                            data = json.loads(data_str)
                            yield data
                        except json.JSONDecodeError:
                            pass
        else:
            # Non-streaming fallback
            response = httpx.post(url, json=payload, headers=headers, timeout=60.0)
            if response.status_code == 200:
                yield response.json()

    def create_embedding(self, input):
        # Support remote embeddings if needed
        url = f"{self.base_url}/v1/embeddings"
        payload = {
            "model": "default",
            "input": input
        }
        try:
            response = httpx.post(url, json=payload, timeout=30.0)
            if response.status_code == 200:
                return response.json()
        except Exception as e:
            logger.error(f"Remote embedding failed: {e}")
        return {'data': [{'embedding': [0.0]*768}]}

class ModelManager:
    def __init__(self):
        self.llms = {}
        self.embed_models = {}
        self.lock = asyncio.Lock()
        import threading
        self.thread_lock = threading.RLock()
        
        # Load Cluster Nodes
        # Format: http://192.168.1.11:8000,http://192.168.1.12:8000
        nodes_env = os.environ.get("CLUSTER_NODES", "")
        self.workers = [n.strip() for n in nodes_env.split(',') if n.strip()]
        if self.workers:
            logger.info(f"AI Cluster Mode Enabled. Workers: {self.workers}")

    def get_llm(self, model_path: str, n_gpu_layers: int = None):
        # 1. Cluster Distribution Logic
        if self.workers:
            # Simple Random Load Balancing
            worker_url = random.choice(self.workers)
            logger.info(f"Delegating inference to worker: {worker_url}")
            return RemoteLlama(worker_url, model_path)

        # 2. Local Logic
        # Thread-safe access for checking cache
        with self.thread_lock:
             # Exclusive Policy: Unload Embed Models if exists to free VRAM
             if self.embed_models:
                 logger.info("Unloading embedding models to free VRAM for LLM...")
                 keys = list(self.embed_models.keys())
                 for k in keys:
                     try:
                         if hasattr(self.embed_models[k], 'close'):
                             self.embed_models[k].close()
                     except: pass
                     del self.embed_models[k]
                 self.embed_models.clear()
                 gc.collect()

             if model_path in self.llms:
                 return self.llms[model_path]
        
             # If not found, load it (using lock to prevent double loading)
             # Double check pattern
             if model_path in self.llms:
                  return self.llms[model_path]
              
             # Unload other models ONLY if it's NOT the discord model or we are loading main model over discord model
             is_discord_model = "discord_" in str(model_path)
             
             if not is_discord_model and self.llms:
                # If loading main model, we might need to unload previous main models.
                # But try to keep discord model if it exists
                logger.info("Cleaning up previous LLMs...")
                keys = list(self.llms.keys())
                for k in keys:
                    if "discord_" in str(k): continue # Keep discord model
                    try:
                        if hasattr(self.llms[k], 'close'):
                            self.llms[k].close()
                        del self.llms[k]
                    except Exception as ex:
                        logger.warning(f"Error closing model {k}: {ex}")
                gc.collect()
            
             logger.info(f"Loading LLM: {model_path}")
             
             # Determine GPU layers based on model size hints
             layers = -1
             if n_gpu_layers is not None:
                 layers = n_gpu_layers
             elif "7b" in model_path.lower() or "8b" in model_path.lower():
                  # For larger models, force CPU to avoid VRAM OOM on limited hardware
                  logger.info("Forcing CPU for large model to ensure stability.")
                  layers = 0 
             
             try:
                 llm = Llama(
                     model_path=model_path, 
                     n_ctx=2048, # Reduced to 2048 to prevent OOM
                     n_batch=64, 
                     n_gpu_layers=layers,
                     verbose=True
                 )
                 self.llms[model_path] = llm
                 return llm
             except Exception as e:
                 logger.error(f"Failed to load LLM {model_path} with GPU: {e}")
                 logger.info("Retrying with CPU fallback...")
                 try:
                     llm = Llama(
                         model_path=model_path, 
                         n_ctx=2048,
                         n_batch=64, 
                         n_gpu_layers=0, # Force CPU
                         verbose=True
                     )
                     self.llms[model_path] = llm
                     return llm
                 except Exception as e2:
                     logger.error(f"Failed to load LLM {model_path} with CPU: {e2}")
                     raise e2

    def get_embed_model(self, model_path: str):
        with self.thread_lock:
            # Exclusive Policy: Unload LLMs if exists to free VRAM for Embedding
            if self.llms:
                logger.info("Unloading LLMs to free VRAM for Embedding...")
                keys = list(self.llms.keys())
                for k in keys:
                    try:
                        if hasattr(self.llms[k], 'close'):
                            self.llms[k].close()
                    except: pass
                    del self.llms[k]
                self.llms.clear()
                gc.collect()

            if model_path in self.embed_models:
                return self.embed_models[model_path]

            logger.info(f"Loading Embedding Model: {model_path}")
            try:
                # For embeddings, we prefer local processing for speed if possible,
                # unless offloading is strictly required. For now, keep local CPU/GPU mixed.
                embed_model = Llama(
                    model_path=model_path,
                    embedding=True,
                    n_gpu_layers=0, # Use CPU for embeddings to save VRAM for chat
                    n_ctx=2048, # 8192 is too large for embeddings alongside chat model and causes OOM crashes
                    verbose=False
                )
                self.embed_models[model_path] = embed_model
                return embed_model
            except Exception as e:
                logger.error(f"Failed to load embedding model: {e}")
                raise e

model_manager = ModelManager()

# --- Database Setup (SQLite) ---
def init_db():
    conn = sqlite3.connect(DB_PATH, timeout=60, check_same_thread=False)
    conn.execute("PRAGMA journal_mode=WAL;")
    cursor = conn.cursor()
    
    # Users Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS users (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        username TEXT UNIQUE NOT NULL,
        display_name TEXT,
        password_hash TEXT NOT NULL,
        role TEXT NOT NULL,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
    )
    ''')
    
    # Migration for display_name if not exists
    try:
        cursor.execute('ALTER TABLE users ADD COLUMN display_name TEXT')
    except sqlite3.OperationalError:
        pass # Already exists
    
    # Settings Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS settings (
        key TEXT PRIMARY KEY,
        value TEXT
    )
    ''')
    
    # Chat Sessions Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS chat_sessions (
        id TEXT PRIMARY KEY,
        user_id INTEGER NOT NULL,
        title TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users (id)
    )
    ''')

    # Chat Messages Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS chat_messages (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        session_id TEXT NOT NULL,
        role TEXT NOT NULL,
        content TEXT NOT NULL,
        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (session_id) REFERENCES chat_sessions (id) ON DELETE CASCADE
    )
    ''')

    # File Index State Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS file_index_state (
        path TEXT PRIMARY KEY,
        modified_time REAL NOT NULL,
        last_seen REAL NOT NULL
    )
    ''')
    
    # User Memory Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS user_memory (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        user_id INTEGER NOT NULL,
        key TEXT,
        value TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (user_id) REFERENCES users (id)
    )
    ''')

    # Conversation Summaries Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS chat_summaries (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        session_id TEXT NOT NULL,
        summary TEXT NOT NULL,
        range_start INTEGER, 
        range_end INTEGER,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (session_id) REFERENCES chat_sessions (id) ON DELETE CASCADE
    )
    ''')
    
    # Canvases Table
    cursor.execute('''
    CREATE TABLE IF NOT EXISTS canvases (
        id TEXT PRIMARY KEY,
        session_id TEXT NOT NULL,
        title TEXT,
        content TEXT,
        language TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        FOREIGN KEY (session_id) REFERENCES chat_sessions (id) ON DELETE CASCADE
    )
    ''')
    
    # Create default users if they don't exist
    cursor.execute('SELECT * FROM users WHERE username = ?', ('adminuser',))
    if not cursor.fetchone():
        hashed = pwd_context.hash('admin')
        cursor.execute('INSERT INTO users (username, display_name, password_hash, role) VALUES (?, ?, ?, ?)', 
                      ('adminuser', 'Administrator', hashed, 'admin'))
        logger.info("Created default admin user: adminuser / admin")
        
    cursor.execute('SELECT * FROM users WHERE username = ?', ('user',))
    if not cursor.fetchone():
        hashed = pwd_context.hash('admin')
        cursor.execute('INSERT INTO users (username, display_name, password_hash, role) VALUES (?, ?, ?, ?)', 
                      ('user', 'General User', hashed, 'user'))
        logger.info("Created default general user: user / admin")
    
    conn.commit()
    conn.close()

# --- Auth Utilities ---
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=status.HTTP_401_UNAUTHORIZED,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
    except JWTError:
        raise credentials_exception
        
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('SELECT id, username, role FROM users WHERE username = ?', (username,))
    user = cursor.fetchone()
    conn.close()
    
    if user is None:
        raise credentials_exception
    return {"id": user[0], "username": user[1], "role": user[2]}

async def get_current_admin(current_user: dict = Depends(get_current_user)):
    if current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Not authorized")
    return current_user

# --- Models ---
class UserCreate(BaseModel):
    username: str
    display_name: str
    password: str
    role: str = "user"

class UserUpdate(BaseModel):
    display_name: str
    password: Optional[str] = None
    role: str

class UserResponse(BaseModel):
    id: int
    username: str
    display_name: Optional[str] = None
    role: str
    created_at: Optional[str] = None

class Token(BaseModel):
    access_token: str
    token_type: str

class ChatRequest(BaseModel):
    message: str
    model_id: str
    use_nas: bool = False
    session_id: Optional[str] = None
    attached_file_ids: Optional[List[str]] = None
    canvas_mode: bool = False

class ChatSession(BaseModel):
    id: str
    title: str
    updated_at: str

class ChatMessageDB(BaseModel):
    id: int
    role: str
    content: str
    timestamp: str

class RenameRequest(BaseModel):
    title: str

class IndexedDocument(BaseModel):
    id: str
    filename: str
    path: str
    chunk_count: int
    modified_at: str

class ChunkSearchRequest(BaseModel):
    query: str = ""

class MemoryEntry(BaseModel):
    key: str
    value: str

class GreetRequest(BaseModel):
    memories: List[MemoryEntry]
    time_of_day: str # "morning", "afternoon", "evening", "night"
    limit: int = 10
    file_path: Optional[str] = None

class ChunkResult(BaseModel):
    id: str
    content: str
    metadata: Dict[str, Any]
    score: Optional[float] = None

class Canvas(BaseModel):
    id: str
    session_id: str
    title: str = "Untitled"
    content: str | None = None
    language: str | None = None
    updated_at: str

class CanvasCreate(BaseModel):
    session_id: str
    title: str = "Untitled Canvas"
    content: str = ""
    language: str = "text"

class CanvasUpdate(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    language: Optional[str] = None

# --- RAG & Indexing ---
def get_chroma_client():
    return chromadb.PersistentClient(path=str(CHROMA_DB_DIR))

class GGUFEmbeddingFunction:
    def __init__(self, model_path):
        self.model_path = model_path
        
    def __call__(self, input: List[str]) -> List[List[float]]:
        # Use the global lock to prevent concurrent GPU/CPU usage during inference
        with model_manager.thread_lock:
            try:
                llm = model_manager.get_embed_model(self.model_path)
                if not llm:
                    logger.error("Embedding model not loaded")
                    return [[] for _ in input] # Return empty if failed
                    
                embeddings = []
                for i, text in enumerate(input):
                    try:
                        # Llama.cpp embedding
                        embed = llm.create_embedding(text)
                        embeddings.append(embed['data'][0]['embedding'])
                    except Exception as e:
                        logger.error(f"Failed to create embedding for text {i}: {e}")
                        # Return zero vector as fallback
                        embeddings.append([0.0] * 768)  # nomic-embed has 768 dimensions
                return embeddings
            except Exception as e:
                logger.error(f"Critical error in embedding function: {e}")
                return [[0.0] * 768 for _ in input]

def read_docx_file(path: Path) -> str:
    if not docx: return ""
    try:
        doc = docx.Document(path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.warning(f"Error reading docx {path}: {e}")
        return ""

def read_excel_file(path: Path) -> str:
    if not openpyxl: return ""
    try:
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
        text = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                text.append(" ".join([str(cell) for cell in row if cell is not None]))
        return "\n".join(text)
    except Exception as e:
        logger.warning(f"Error reading xlsx {path}: {e}")
        return ""

def read_pdf_file(path: Path) -> str:
    if not pypdf: return ""
    try:
        reader = pypdf.PdfReader(str(path))
        text_content = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text_content.append(t)
        return "\n".join(text_content)
    except Exception as e:
        logger.warning(f"Error reading pdf {path}: {e}")
        return ""

def recursive_character_text_splitter(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """
    Splits text recursively by separators to keep related content together.
    Separators: \n\n, \n, . , space, empty
    """
    separators = ["\n\n", "\n", ". ", " ", ""]
    
    def _split_text(text: str, separators: List[str]) -> List[str]:
        final_chunks = []
        separator = separators[-1]
        new_separators = []
        
        for i, sep in enumerate(separators):
            if sep == "":
                separator = ""
                break
            if sep in text:
                separator = sep
                new_separators = separators[i+1:]
                break
        
        splits = text.split(separator) if separator else list(text)
        
        good_splits = []
        for s in splits:
            if s.strip():
                good_splits.append(s if separator == "" else s + separator)
        
        current_chunk = ""
        for s in good_splits:
            if len(current_chunk) + len(s) < chunk_size:
                current_chunk += s
            else:
                if current_chunk:
                    final_chunks.append(current_chunk)
                current_chunk = s
                if len(current_chunk) > chunk_size and new_separators:
                    # Recursively split if still too big
                    sub_chunks = _split_text(current_chunk, new_separators)
                    final_chunks.extend(sub_chunks)
                    current_chunk = ""
        
        if current_chunk:
            final_chunks.append(current_chunk)
            
        return final_chunks

    return _split_text(text, separators)

import psutil
try:
    import psutil
except ImportError:
    psutil = None

def index_documents_task():
    """Background task for indexing with DB-backed state and stop capability"""
    state.is_indexing = True
    state.stop_indexing_flag = False
    state.indexing_status = "Starting..."
    state.indexing_progress = 0
    state.indexing_log = []
    state.indexing_start_time = datetime.now()
    state.indexing_processed_files = 0
    state.indexing_total_files = 0
    
    process = psutil.Process(os.getpid()) if psutil else None

    def log(msg):
        timestamp = datetime.now().strftime("%H:%M:%S")
        mem_usage = f"{process.memory_info().rss / 1024 / 1024:.2f} MB" if process else "N/A"
        entry = f"[{timestamp}][{mem_usage}] {msg}"
        state.indexing_log.append(entry)
        if len(state.indexing_log) > 2000:
            state.indexing_log.pop(0)
        logger.info(entry)

    log("Starting indexing task...")
    
    db_conn = None
    try:
        # --- Setup Phase ---
        scan_start_time = time.time()
        
        source_dir = MNT_DIR if state.current_storage_mode == "nas" else INTERNAL_NAS_DIR
        log(f"Source directory: {source_dir}")

        db_conn = sqlite3.connect(DB_PATH, timeout=60, check_same_thread=False)
        db_cursor = db_conn.cursor()
        log("Database connection established.")

        embed_model_name = "nomic-embed-text-v1.5.f16.gguf"
        embed_model_path = MODELS_DIR / embed_model_name
        if not embed_model_path.exists():
            log(f"CRITICAL: Embedding model {embed_model_name} missing.")
            state.indexing_status = "Error: Model missing"
            return
        
        log("Initializing embedding function...")
        embedding_fn = GGUFEmbeddingFunction(str(embed_model_path))
        log("Embedding function initialized.")
        
        log("Initializing ChromaDB...")
        client = get_chroma_client()
        collection = client.get_or_create_collection(name="nas_documents", embedding_function=embedding_fn)
        log("ChromaDB collection loaded.")

        # --- Scanning & Indexing Phase ---
        log("Starting scan and index process...")
        
        batch_size = 10
        current_batch_ids, current_batch_docs, current_batch_metadatas = [], [], []
        scanned_count, processed_count = 0, 0
        
        for root, _, files in os.walk(source_dir):
            if state.stop_indexing_flag:
                log("Stop flag received, breaking scan loop.")
                break
                
            for file in files:
                if state.stop_indexing_flag: break
                
                scanned_count += 1
                file_path = Path(root) / file
                file_key = str(file_path)
                
                if scanned_count % 100 == 0:
                    state.indexing_status = f"Scanned {scanned_count}, Processed {processed_count}..."
                    db_conn.commit()  # Periodically commit updates


                if not file.endswith(('.txt', '.md', '.json', '.py', '.js', '.ts', '.html', '.css', '.csv', '.docx', '.xlsx')):
                    continue
                
                log(f"[{scanned_count}] Checking: {file_key}")
                    
                try:
                    stat = file_path.stat()
                    mod_time = stat.st_mtime

                    if stat.st_size > 1024 * 1024 * 1024:
                        continue

                    db_cursor.execute("SELECT modified_time FROM file_index_state WHERE path = ?", (file_key,))
                    result = db_cursor.fetchone()
                    
                    if result and result[0] == mod_time:
                        db_cursor.execute("UPDATE file_index_state SET last_seen = ? WHERE path = ?", (scan_start_time, file_key))
                        continue
                        
                    log(f"  -> Processing required for: {file_key}")
                    state.indexing_status = f"Indexing: {file}..."
                    
                    content = ""
                    log(f"    - Reading content...")
                    if file_path.suffix == '.docx':
                        content = read_docx_file(file_path)
                    elif file_path.suffix == '.xlsx':
                        content = read_excel_file(file_path)
                    else:
                        if stat.st_size > 10 * 1024 * 1024:
                            log(f"    - Skipping text read >10MB for: {file}")
                            continue
                        try:
                            content = file_path.read_text(encoding='utf-8', errors='ignore')
                        except Exception as read_err:
                            log(f"    - Read error: {read_err}")
                            continue
                    log(f"    - Content read. Length: {len(content)}")

                    if not content or not content.strip():
                        db_cursor.execute("INSERT OR REPLACE INTO file_index_state (path, modified_time, last_seen) VALUES (?, ?, ?)",
                                          (file_key, mod_time, scan_start_time))
                        log(f"    - Empty content. Skipping.")
                        continue

                    log(f"    - Chunking content...")
                    chunks = recursive_character_text_splitter(content, chunk_size=1000, chunk_overlap=200)
                    log(f"    - Content chunked into {len(chunks)} parts.")
                    
                    file_hash = hashlib.md5(file_key.encode()).hexdigest()
                    mod_time_iso = datetime.fromtimestamp(mod_time).isoformat()

                    log(f"    - Deleting old chunks from ChromaDB...")
                    collection.delete(where={"path": file_key})
                    log(f"    - Old chunks deleted.")

                    for j, chunk in enumerate(chunks):
                        chunk_id = f"{file_hash}_{j}"
                        current_batch_ids.append(chunk_id)
                        current_batch_docs.append(chunk)
                        current_batch_metadatas.append({"filename": file_path.name, "path": file_key, "modified_at": mod_time_iso, "chunk_index": j, "total_chunks": len(chunks)})
                        
                        # Process batch if it reaches batch_size, even within a single file
                        if len(current_batch_ids) >= batch_size:
                            log(f"    - Adding batch of {len(current_batch_ids)} chunks to ChromaDB...")
                            try:
                                collection.add(ids=current_batch_ids, documents=current_batch_docs, metadatas=current_batch_metadatas)
                                log(f"    - Batch added successfully.")
                            except Exception as add_err:
                                log(f"    - ERROR adding batch to ChromaDB: {add_err}")
                                # Continue processing despite error
                            current_batch_ids, current_batch_docs, current_batch_metadatas = [], [], []

                    db_cursor.execute("INSERT OR REPLACE INTO file_index_state (path, modified_time, last_seen) VALUES (?, ?, ?)",
                                      (file_key, mod_time, scan_start_time))
                    db_conn.commit()
                    
                    processed_count += 1
                    state.indexing_processed_files = processed_count
                    state.indexing_total_files = scanned_count

                except Exception as e:
                    log(f"  -> CRITICAL ERROR processing {file}: {e}")

        # Final Batch
        if current_batch_ids and not state.stop_indexing_flag:
            log(f"Adding final batch of {len(current_batch_ids)} chunks...")
            try:
                collection.add(ids=current_batch_ids, documents=current_batch_docs, metadatas=current_batch_metadatas)
                db_conn.commit()
                log("Final batch added and DB committed.")
            except Exception as e:
                log(f"ERROR adding final batch: {e}")

        # --- Deletion Handling Phase ---
        if not state.stop_indexing_flag:
            log("Checking for deleted files...")
            db_cursor.execute("SELECT path FROM file_index_state WHERE last_seen < ?", (scan_start_time,))
            deleted_files = db_cursor.fetchall()
            
            if deleted_files:
                deleted_paths = [row[0] for row in deleted_files]
                log(f"Found {len(deleted_paths)} deleted files. Removing from index...")
                
                for i in range(0, len(deleted_paths), 100):
                    batch_paths = deleted_paths[i:i+100]
                    log(f"  - Deleting batch of {len(batch_paths)} from ChromaDB.")
                    collection.delete(where={"path": {"$in": batch_paths}})
                
                log("  - Deleting records from state DB.")
                db_cursor.execute("DELETE FROM file_index_state WHERE last_seen < ?", (scan_start_time,))
                db_conn.commit()
                log("Finished removing deleted files.")

        if state.stop_indexing_flag:
            state.indexing_status = "Stopped"
            log("Indexing process stopped by user.")
        else:
            state.indexing_status = "Completed"
            state.last_indexed_at = datetime.now().isoformat()
            log(f"Indexing completed successfully. Scanned {scanned_count} files.")
            db_cursor.execute('INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)', ('last_indexed_at', state.last_indexed_at))
            db_conn.commit()
            
    except Exception as e:
        log(f"--- !!! CRITICAL INDEXING FAILURE !!! --- : {e}")
        state.indexing_status = f"Failed: {str(e)}"
        if db_conn:
            db_conn.rollback()
    finally:
        log("Indexing task finished. Closing DB connection.")
        if db_conn:
            db_conn.close()
        state.is_indexing = False

# Discord Agent Logic Removed


# --- FastAPI App ---
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    init_db()

    # Reset stuck indexing state if present
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'indexing_status'")
        row = cursor.fetchone()
        if row:
            status = json.loads(row[0])
            if status.get("is_indexing"):
                logger.warning("Found stuck indexing state on startup. Resetting to Idle.")
                status["is_indexing"] = False
                status["status"] = "Interrupted"
                cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", 
                              ("indexing_status", json.dumps(status)))
                conn.commit()
        conn.close()
    except Exception as e:
        logger.error(f"Failed to reset indexing state: {e}")
    
    # Load settings
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('SELECT value FROM settings WHERE key = ?', ('storage_mode',))
    row = cursor.fetchone()
    if row:
        state.current_storage_mode = row[0]
        
    cursor.execute('SELECT value FROM settings WHERE key = ?', ('last_indexed_at',))
    row = cursor.fetchone()
    if row:
        state.last_indexed_at = row[0]
    
    conn.close()

    # Ensure model directories for all users
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('SELECT username FROM users')
    all_users = cursor.fetchall()
    conn.close()
    
    for (u_name,) in all_users:
        ensure_user_models_dir(u_name)

    # Preload Fast model for adminuser as a warm-up (optional)
    admin_models = ensure_user_models_dir("adminuser")
    fast_model_path = admin_models / "qwen2-1.5b-instruct-q8_0.gguf"
    if fast_model_path.exists():
        logger.info(f"Preloading Fast model: {fast_model_path}")
        try:
            model_manager.get_llm(str(fast_model_path))
            logger.info("Fast model preloaded successfully.")
        except Exception as e:
            logger.error(f"Failed to preload Fast model: {e}")
    else:
        logger.warning(f"Fast model not found at {fast_model_path}, skipping preload.")
    
    
    yield
    
    # Shutdown


app = FastAPI(lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- Endpoints ---

# --- Canvas Endpoints ---

@app.get("/api/canvases", response_model=List[Canvas])
async def list_canvases(current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    # Join with sessions to ensure user owns the session
    cursor.execute('''
        SELECT c.* FROM canvases c
        JOIN chat_sessions s ON c.session_id = s.id
        WHERE s.user_id = ?
        ORDER BY c.updated_at DESC
    ''', (current_user['id'],))
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]

@app.get("/api/chat/sessions/{session_id}/canvases", response_model=List[Canvas])
async def list_session_canvases(session_id: str, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    # Verify ownership
    cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (session_id,))
    session = cursor.fetchone()
    if not session or session['user_id'] != current_user['id']:
        conn.close()
        raise HTTPException(status_code=404, detail="Session not found")

    cursor.execute('SELECT * FROM canvases WHERE session_id = ? ORDER BY updated_at DESC', (session_id,))
    rows = cursor.fetchall()
    conn.close()
    return [dict(row) for row in rows]

@app.post("/api/canvases", response_model=Canvas)
async def create_canvas(data: CanvasCreate, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    # Logic to handle auto-creation of Docs Storage Session
    docs_session_id = f"docs_storage_{current_user['id']}"
    
    # Remap placeholders from frontend
    if data.session_id in ['docs_create_session', 'temp_doc_init', 'docs_storage', 'docs_storage_auto']:
        data.session_id = docs_session_id

    # Verify session ownership
    cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (data.session_id,))
    session = cursor.fetchone()
    
    if not session:
        # If it is the special docs session, create it on demand
        if data.session_id == docs_session_id:
            now_sess = datetime.utcnow().isoformat()
            cursor.execute('''
                INSERT INTO chat_sessions (id, user_id, title, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?)
            ''', (docs_session_id, current_user['id'], 'Docs Storage', now_sess, now_sess))
            conn.commit() 
            session = {'user_id': current_user['id']}
        else:
            conn.close()
            raise HTTPException(status_code=404, detail="Session not found")
            
    if session['user_id'] != current_user['id']:
        conn.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    canvas_id = str(uuid.uuid4())
    now = datetime.utcnow().isoformat()
    
    cursor.execute('''
        INSERT INTO canvases (id, session_id, title, content, language, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    ''', (canvas_id, data.session_id, data.title, data.content, data.language, now, now))
    
    conn.commit()
    conn.close()
    
    return {
        "id": canvas_id,
        "session_id": data.session_id,
        "title": data.title,
        "content": data.content,
        "language": data.language,
        "created_at": now,
        "updated_at": now
    }

@app.put("/api/canvases/{canvas_id}")
async def update_canvas(canvas_id: str, data: CanvasUpdate, session_id: Optional[str] = Body(None), current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM canvases WHERE id = ?', (canvas_id,))
    existing = cursor.fetchone()
    
    if existing:
        # Check ownership via session
        cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (existing['session_id'],))
        sess = cursor.fetchone()
        if not sess or sess['user_id'] != current_user['id']:
            conn.close()
            raise HTTPException(status_code=403, detail="Not authorized")
            
        update_fields = []
        params = []
        if data.title is not None:
            update_fields.append("title = ?")
            params.append(data.title)
        if data.content is not None:
            update_fields.append("content = ?")
            params.append(data.content)
        if data.language is not None:
            update_fields.append("language = ?")
            params.append(data.language)
            
        update_fields.append("updated_at = CURRENT_TIMESTAMP")
        
        if update_fields:
            sql = f"UPDATE canvases SET {', '.join(update_fields)} WHERE id = ?"
            params.append(canvas_id)
            cursor.execute(sql, tuple(params))
            conn.commit()
            
        # Return updated or created
        cursor.execute('SELECT * FROM canvases WHERE id = ?', (canvas_id,))
        updated = cursor.fetchone()
        conn.close()
        return dict(updated)
    else:
        # Create new
        if not session_id:
             conn.close()
             raise HTTPException(status_code=400, detail="session_id required for new canvas")
             
        # Verify session ownership
        cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (session_id,))
        sess = cursor.fetchone()
        if not sess or sess['user_id'] != current_user['id']:
            conn.close()
            raise HTTPException(status_code=404, detail="Session not found")
            
        cursor.execute('''
            INSERT INTO canvases (id, session_id, title, content, language)
            VALUES (?, ?, ?, ?, ?)
        ''', (canvas_id, session_id, data.title or "Untitled", data.content or "", data.language or "text"))
        conn.commit()
        
    conn.close()
    return {"status": "ok"}

@app.delete("/api/canvases/{canvas_id}")
async def delete_canvas(canvas_id: str, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Check ownership
    cursor.execute('''
        SELECT s.user_id FROM canvases c
        JOIN chat_sessions s ON c.session_id = s.id
        WHERE c.id = ?
    ''', (canvas_id,))
    row = cursor.fetchone()
    
    if not row or row[0] != current_user['id']:
        conn.close()
        raise HTTPException(status_code=404, detail="Canvas not found")
        
    cursor.execute('DELETE FROM canvases WHERE id = ?', (canvas_id,))
    conn.commit()
    conn.close()
    return {"status": "deleted"}

@app.post("/token", response_model=Token)
async def login_for_access_token(form_data: OAuth2PasswordRequestForm = Depends()):
    try:
        conn = sqlite3.connect(DB_PATH, timeout=60, check_same_thread=False)
        cursor = conn.cursor()
        cursor.execute('SELECT username, password_hash, role FROM users WHERE username = ?', (form_data.username,))
        user = cursor.fetchone()
        conn.close()
        
        if not user or not verify_password(form_data.password, user[1]):
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Incorrect username or password",
                headers={"WWW-Authenticate": "Bearer"},
            )
        
        access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
        access_token = create_access_token(
            data={"sub": user[0], "role": user[2]}, expires_delta=access_token_expires
        )
        return {"access_token": access_token, "token_type": "bearer"}
    except Exception as e:
        with open("login_error.log", "w") as f:
            import traceback
            f.write(traceback.format_exc())
        raise e

@app.get("/api/users/me")
async def read_users_me(current_user: dict = Depends(get_current_user)):
    return current_user

@app.get("/api/users/me/memory")
async def get_user_memory(current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("SELECT key, value FROM user_memory WHERE user_id = ?", (current_user["id"],))
    memories = [{"key": r[0], "value": r[1]} for r in cursor.fetchall()]
    conn.close()
    return memories

@app.post("/api/users/me/memory")
async def save_user_memory(entry: MemoryEntry, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO user_memory (user_id, key, value) VALUES (?, ?, ?)", 
                   (current_user["id"], entry.key, entry.value))
    conn.commit()
    conn.close()
    return {"status": "saved"}

@app.post("/api/ai/greet")
async def ai_greet(req: GreetRequest, current_user: dict = Depends(get_current_user)):
    """Generate a personalized greeting using the AI"""
    try:
        # Build memory context
        mem_str = "\n".join([f"- {m.key}: {m.value}" for m in req.memories])
        
        system_prompt = (
            "あなたはユーザーに寄り添う、最高にスマートで美しいAI秘書です。\n"
            "以下の【ユーザーの記憶】と【現在の時間帯】を元に、極限までシンプルかつ心に響く、最高に美しい出迎えの一言を日本語で作成してください。\n"
            "【制約】\n"
            "- 返信は一言（1〜2文）のみ。\n"
            "- 「お疲れ様です」や「おはようございます」などの挨拶を含める。\n"
             "- 敬語（丁寧語）を使う。\n"
            "- 余計な解説は不要。挨拶文のみを出力せよ。"
        )
        
        user_prompt = f"【現在の時間帯】: {req.time_of_day}\n【ユーザーの記憶】:\n{mem_str or 'まだ記憶はありません。'}\n\n出迎えのメッセージ:"
        
        # Call model (Fast model)
        # Find qwen2 1.5b or 3b
        model_path = None
        for p, llm in model_manager.llms.items():
            if "qwen2" in str(p).lower():
                model_path = p
                break
        
        if not model_path:
            # Fallback to defaults or first loaded
            if model_manager.llms:
                model_path = list(model_manager.llms.keys())[0]
        
        if not model_path:
             return {"greeting": f"{req.time_of_day}。今日も素晴らしい一日になりますように。"}

        llm = model_manager.get_llm(model_path)
        res = llm.create_chat_completion(
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=64,
            temperature=0.8
        )
        
        greeting = res['choices'][0]['message']['content'].strip()
        return {"greeting": greeting}
    except Exception as e:
        logger.error(f"Greet Error: {e}")
        return {"greeting": "お疲れ様です。今日もあなたの創造性を最大限に引き出しましょう。"}

# Admin Endpoints
@app.get("/api/admin/users", response_model=List[UserResponse])
async def get_users(admin: dict = Depends(get_current_admin)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('SELECT id, username, display_name, role, created_at FROM users')
    users = [{"id": r[0], "username": r[1], "display_name": r[2], "role": r[3], "created_at": r[4]} for r in cursor.fetchall()]
    conn.close()
    return users

@app.post("/api/admin/users", response_model=UserResponse)
async def create_user(user: UserCreate, background_tasks: BackgroundTasks, admin: dict = Depends(get_current_admin)):
    # Validation
    import re
    if not re.match(r'^[a-zA-Z0-9-]+$', user.username):
        raise HTTPException(status_code=400, detail="Username must be alphanumeric or hyphen only")
    if '_' in user.password or '.' in user.password:
        raise HTTPException(status_code=400, detail="Password cannot contain '_' or '.'")

    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    try:
        hashed = get_password_hash(user.password)
        cursor.execute('INSERT INTO users (username, display_name, password_hash, role) VALUES (?, ?, ?, ?)', 
                      (user.username, user.display_name, hashed, user.role))
        user_id = cursor.lastrowid
        conn.commit()
        
        # Trigger model directory creation in background
        background_tasks.add_task(ensure_user_models_dir, user.username)
        
        return {"id": user_id, "username": user.username, "display_name": user.display_name, "role": user.role, "created_at": datetime.now().isoformat()}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Username already registered")
    finally:
        conn.close()

@app.put("/api/admin/users/{user_id}", response_model=UserResponse)
async def update_user(user_id: int, user: UserUpdate, admin: dict = Depends(get_current_admin)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Check if user exists
    cursor.execute('SELECT username FROM users WHERE id = ?', (user_id,))
    row = cursor.fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")
    
    username = row[0]
    
    # Validation
    if user.password and ('_' in user.password or '.' in user.password):
        conn.close()
        raise HTTPException(status_code=400, detail="Password cannot contain '_' or '.'")

    try:
        if user.password:
            hashed = get_password_hash(user.password)
            cursor.execute('UPDATE users SET display_name = ?, password_hash = ?, role = ? WHERE id = ?', 
                          (user.display_name, hashed, user.role, user_id))
        else:
            cursor.execute('UPDATE users SET display_name = ?, role = ? WHERE id = ?', 
                          (user.display_name, user.role, user_id))
        conn.commit()
        return {"id": user_id, "username": username, "display_name": user.display_name, "role": user.role, "created_at": datetime.now().isoformat()}
    finally:
        conn.close()

@app.delete("/api/admin/users/{user_id}")
async def delete_user(user_id: int, admin: dict = Depends(get_current_admin)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Check if user is adminuser
    cursor.execute('SELECT username FROM users WHERE id = ?', (user_id,))
    row = cursor.fetchone()
    if row and row[0] == 'adminuser':
        conn.close()
        raise HTTPException(status_code=400, detail="Cannot delete the default admin user")
        
    cursor.execute('DELETE FROM users WHERE id = ?', (user_id,))
    conn.commit()
    conn.close()
    return {"status": "success"}

@app.get("/api/admin/nas/status")
async def get_nas_status(admin: dict = Depends(get_current_admin)):
    # Check if MNT_DIR is mounted or has content
    is_mounted = os.path.ismount(MNT_DIR)
    # Also check if it has files (sometimes manual mount might not show as ismount in container/some envs)
    has_files = False
    try:
        if any(os.scandir(MNT_DIR)):
            has_files = True
    except:
        pass
        
    status = get_db_status()
    
    # Log from DB status (real-time from indexer)
    log_content = status.get("indexing_log", [])
    
    # Fallback/Merge with file log if needed, or just use DB log
    if not log_content:
        log_path = Path("logs/indexing.log")
        if log_path.exists():
            try:
                with open(log_path, "r") as f:
                    lines = f.readlines()
                    log_content = [l.strip() for l in lines[-50:]]
            except Exception:
                pass


    # Get total indexed documents count from DB
    total_indexed_documents = 0
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT COUNT(*) FROM file_index_state")
        total_indexed_documents = cursor.fetchone()[0]
        
        # Get correct last_indexed_at
        cursor.execute("SELECT value FROM settings WHERE key = 'last_indexed_at'")
        last_indexed_row = cursor.fetchone()
        last_indexed_at = last_indexed_row[0] if last_indexed_row else None
        conn.close()
    except Exception as e:
        logger.error(f"Error fetching stats: {e}")
        last_indexed_at = None

    # Calculate ChromaDB usage
    chroma_usage = 0
    try:
        if CHROMA_DB_DIR.exists():
            chroma_usage = sum(f.stat().st_size for f in CHROMA_DB_DIR.glob('**/*') if f.is_file())
    except Exception as e:
        logger.error(f"Error calculating chroma usage: {e}")

    return {
        "is_mounted": is_mounted or has_files,
        "mount_path": str(MNT_DIR),
        "storage_mode": get_storage_mode(),
        "is_indexing": status.get("is_indexing", False),
        "indexing_progress": status.get("progress", 0),
        "indexing_status": status.get("status", "Idle"),
        "indexing_log": [l.strip() for l in log_content],
        "total_files": status.get("total_files", 0),
        "processed_files": status.get("processed_files", 0),
        "last_indexed_at": last_indexed_at,
        "total_indexed_documents": total_indexed_documents,
        "chroma_usage": chroma_usage
    }

@app.post("/api/admin/nas/mode")
async def set_storage_mode(mode: str = Body(..., embed=True), admin: dict = Depends(get_current_admin)):
    if mode not in ["nas", "internal"]:
        raise HTTPException(status_code=400, detail="Invalid mode")
    
    state.current_storage_mode = mode
    
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)', ('storage_mode', mode))
    conn.commit()
    conn.close()
    
    return {"status": "success", "mode": mode}

@app.post("/api/admin/index")
async def trigger_indexing(background_tasks: BackgroundTasks, storage_mode: str = Body('nas', embed=True), admin: dict = Depends(get_current_admin)):
    status = get_db_status()
    if status.get("is_indexing"):
        raise HTTPException(status_code=400, detail="Indexing already in progress")
    
    logger.info("Triggering indexing process...")
    # Run indexer.py in a separate process
    subprocess.Popen([sys.executable, "indexer.py", storage_mode])
    
    return {"status": "started", "storage_mode": storage_mode}

@app.post("/api/admin/index/stop")
async def stop_indexing(admin: dict = Depends(get_current_admin)):
    try:
        # Immediately signal the running task to stop in memory
        state.stop_indexing_flag = True
        
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("stop_indexing_flag", "true"))
        conn.commit()
        conn.close()
        return {"status": "stopping"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/admin/index/clear")
async def clear_indexing_status(admin: dict = Depends(get_current_admin)):
    try:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        
        # 1. Clear SQLite Status
        # Also clear in-memory state to reflect immediately
        state.indexing_log = []
        state.indexing_processed_files = 0
        state.indexing_total_files = 0
        
        empty_status = {
            "status": "Idle",
            "progress": 0,
            "is_indexing": False,
            "processed_files": 0,
            "total_files": 0,
            "last_updated": datetime.now().isoformat(),
            "indexing_log": []
        }
        cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", 
                      ("indexing_status", json.dumps(empty_status)))
                      
        # 2. Clear File Index State (to force re-scan)
        cursor.execute("DELETE FROM file_index_state")
        cursor.execute("DELETE FROM settings WHERE key = 'last_indexed_at'")
        conn.commit()
        conn.close()

        # 3. Clear ChromaDB Collections
        try:
            client = get_chroma_client()
            for col_name in ["documents_nas", "documents_internal"]:
                try:
                    client.delete_collection(col_name)
                    logger.info(f"Deleted collection: {col_name}")
                except Exception as e:
                    logger.warning(f"Failed to delete collection {col_name} (may not exist): {e}")
        except Exception as chroma_err:
             logger.error(f"Failed to clear ChromaDB: {chroma_err}")

        return {"status": "cleared"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))



# --- Agent Integration ---
try:
    from agent_core import AgentGateway
    agent_gateway = AgentGateway(model_manager)
except ImportError as e:
    logger.error(f"Failed to import Agent Core: {e}")
    agent_gateway = None

@app.post("/api/admin/agent/config")
async def configure_agent(data: dict = Body(...), admin: dict = Depends(get_current_admin)):
    """Configure the models for the agent"""
    reflex_model = data.get("reflex_model")
    planner_model = data.get("planner_model")
    
    if not reflex_model or not planner_model:
        raise HTTPException(status_code=400, detail="Both reflex_model and planner_model are required")
        
    # Verify models exist in MODELS_DIR
    r_path = MODELS_DIR / reflex_model
    p_path = MODELS_DIR / planner_model
    
    if not r_path.exists() or not p_path.exists():
         raise HTTPException(status_code=400, detail="One or both model files not found")

    # Save to DB settings
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("agent_reflex_model", reflex_model))
    cursor.execute("INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)", ("agent_planner_model", planner_model))
    conn.commit()
    conn.close()
    
    # Initialize Gateway
    if agent_gateway:
        agent_gateway.initialize(str(r_path), str(p_path))
        
    return {"status": "configured", "reflex": reflex_model, "planner": planner_model}

@app.post("/api/agent/chat")
async def agent_chat(request: ChatRequest, current_user: dict = Depends(get_current_user)):
    """Unified entry point for talking to the Oonanji Agent"""
    if not agent_gateway:
        raise HTTPException(status_code=503, detail="Agent system not available")
        
    # Auto-initialize if needed (from cold start)
    if not agent_gateway.agent:
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute("SELECT value FROM settings WHERE key = 'agent_reflex_model'")
        r = cursor.fetchone()
        cursor.execute("SELECT value FROM settings WHERE key = 'agent_planner_model'")
        p = cursor.fetchone()
        conn.close()
        
        if r and p:
            r_path = MODELS_DIR / r[0]
            p_path = MODELS_DIR / p[0]
            if r_path.exists() and p_path.exists():
                agent_gateway.initialize(str(r_path), str(p_path))
            else:
                 return {"role": "assistant", "content": "Agent Configuration Error: Model files missing. Please reconfigure in Admin."}
        else:
            return {"role": "assistant", "content": "Agent not configured. Please ask Admin to set Reflex and Planner models."}

    # Process Request
    session_id = request.session_id or str(uuid.uuid4())
    
    try:
        response_text = await agent_gateway.handle_request(session_id, request.message)
        
        # Save to DB (optional, but good for persistence)
        # Note: We should reuse the chat_messages logic here eventually
        
        return {
            "role": "assistant", 
            "content": response_text,
            "session_id": session_id
        }
    except Exception as e:
        logger.error(f"Agent Error: {e}")
        return {"role": "assistant", "content": f"I crashed: {e}", "session_id": session_id}

# --- Model Management APIs ---
MODEL_DOWNLOAD_TASKS = {} # task_id -> {status, progress, total, filename, error}

class ModelDownloadRequest(BaseModel):
    url: str
    filename: str

async def download_model_background(task_id: str, url: str, filename: str):
    MODEL_DOWNLOAD_TASKS[task_id] = {
        "status": "downloading",
        "progress": 0,
        "total": 0,
        "filename": filename
    }
    
    target_path = MODELS_DIR / filename
    temp_path = target_path.with_suffix(".tmp")
    
    try:
        async with httpx.AsyncClient(timeout=None) as client:
            async with client.stream("GET", url, follow_redirects=True) as resp:
                if resp.status_code != 200:
                    raise Exception(f"HTTP {resp.status_code}")
                
                total = int(resp.headers.get("content-length", 0))
                MODEL_DOWNLOAD_TASKS[task_id]["total"] = total
                
                downloaded = 0
                with open(temp_path, "wb") as f:
                    async for chunk in resp.aiter_bytes(chunk_size=1024*1024): # 1MB chunks
                        f.write(chunk)
                        downloaded += len(chunk)
                        
                        if total > 0:
                            progress = int((downloaded / total) * 100)
                            MODEL_DOWNLOAD_TASKS[task_id]["progress"] = progress
                            
        # Rename on success
        if target_path.exists():
            target_path.unlink() # Overwrite if exists
        temp_path.rename(target_path)
        
        # Sync to all existing user directories
        logger.info(f"Syncing {filename} to all user directories...")
        for user_dir in BASE_DIR.glob("models_*"):
            if user_dir.is_dir():
                # Admin uses base MODELS_DIR, so no need to copy to models_adminuser (which shouldn't exist ideally)
                if user_dir.name == "models_adminuser":
                    continue
                    
                try:
                    user_model_path = user_dir / filename
                    logger.info(f"Copying to {user_dir.name}...")
                    shutil.copy2(target_path, user_model_path)
                except Exception as sync_err:
                    logger.error(f"Failed to sync to {user_dir.name}: {sync_err}")

        MODEL_DOWNLOAD_TASKS[task_id]["status"] = "completed"
        MODEL_DOWNLOAD_TASKS[task_id]["progress"] = 100
        
    except Exception as e:
        logger.error(f"Download failed: {e}")
        MODEL_DOWNLOAD_TASKS[task_id]["status"] = "error"
        MODEL_DOWNLOAD_TASKS[task_id]["error"] = str(e)
        if temp_path.exists():
            temp_path.unlink()

@app.get("/api/models/list")
async def list_models_api(current_user: dict = Depends(get_current_user)):
    """List available models in the models directory"""
    if not MODELS_DIR.exists():
        return []
    
    models = []
    for f in MODELS_DIR.glob("*.gguf"):
        models.append({
            "name": f.name,
            "size": f.stat().st_size,
            "modified": datetime.fromtimestamp(f.stat().st_mtime).isoformat()
        })
    return models

@app.post("/api/models/download")
async def start_model_download(req: ModelDownloadRequest, background_tasks: BackgroundTasks, current_user: dict = Depends(get_current_user)):
    # Check if already exists
    if (MODELS_DIR / req.filename).exists():
        raise HTTPException(status_code=409, detail="Model already exists")

    task_id = str(uuid.uuid4())
    background_tasks.add_task(download_model_background, task_id, req.url, req.filename)
    return {"task_id": task_id, "status": "started"}

@app.get("/api/models/download/{task_id}")
async def get_download_status(task_id: str, current_user: dict = Depends(get_current_user)):
    return MODEL_DOWNLOAD_TASKS.get(task_id, {"status": "not_found"})

@app.get("/api/admin/index/documents", response_model=List[IndexedDocument])
async def get_indexed_documents(admin: dict = Depends(get_current_admin)):
    try:
        client = get_chroma_client()
        collection = client.get_collection("nas_documents")
        
        # Get all metadata
        # Note: This might be heavy if millions of chunks. 
        # For now, we'll aggregate from metadata.
        result = collection.get(include=['metadatas'])
        metadatas = result['metadatas']
        
        files = {}
        for m in metadatas:
            path = m['path']
            if path not in files:
                files[path] = {
                    "id": hashlib.md5(path.encode()).hexdigest(),
                    "filename": m['filename'],
                    "path": path,
                    "chunk_count": 0,
                    "modified_at": m.get('modified_at', '')
                }
            files[path]['chunk_count'] += 1
            
        return list(files.values())
    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        return []

@app.post("/api/admin/index/search", response_model=List[ChunkResult])
async def search_indexed_chunks(request: ChunkSearchRequest, admin: dict = Depends(get_current_admin)):
    try:
        client = get_chroma_client()
        collection = client.get_collection("nas_documents")
        
        if request.file_path:
            # Filter by specific file
            result = collection.get(
                where={"path": request.file_path},
                limit=request.limit,
                include=['documents', 'metadatas']
            )
            
            chunks = []
            if result['ids']:
                for i, id in enumerate(result['ids']):
                    chunks.append({
                        "id": id,
                        "content": result['documents'][i],
                        "metadata": result['metadatas'][i]
                    })
            return chunks
            
        elif request.query:
            # Semantic search
            embed_model_name = "nomic-embed-text-v1.5.f16.gguf"
            embed_model_path = MODELS_DIR / embed_model_name
            
            if not embed_model_path.exists():
                raise HTTPException(status_code=500, detail="Embedding model missing")
                
            embedding_fn = GGUFEmbeddingFunction(str(embed_model_path))
            query_embed = embedding_fn([request.query])[0]
            
            results = collection.query(
                query_embeddings=[query_embed],
                n_results=request.limit
            )
            
            chunks = []
            if results['ids']:
                for i, id in enumerate(results['ids'][0]):
                    chunks.append({
                        "id": id,
                        "content": results['documents'][0][i],
                        "metadata": results['metadatas'][0][i],
                        "score": results['distances'][0][i] if 'distances' in results else None
                    })
            return chunks
        else:
            return []
            
    except Exception as e:
        logger.error(f"Error searching chunks: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# --- Chat History Endpoints ---

@app.get("/api/chat/sessions", response_model=List[ChatSession])
async def get_chat_sessions(current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT id, title, updated_at FROM chat_sessions 
        WHERE user_id = ? 
        ORDER BY updated_at DESC
    ''', (current_user['id'],))
    sessions = [{"id": r[0], "title": r[1], "updated_at": r[2]} for r in cursor.fetchall()]
    conn.close()
    return sessions

@app.get("/api/chat/sessions/{session_id}", response_model=List[ChatMessageDB])
async def get_session_messages(session_id: str, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Verify ownership
    cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (session_id,))
    row = cursor.fetchone()
    if not row or row[0] != current_user['id']:
        conn.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    cursor.execute('''
        SELECT id, role, content, timestamp FROM chat_messages 
        WHERE session_id = ? 
        ORDER BY id ASC
    ''', (session_id,))
    messages = [{"id": r[0], "role": r[1], "content": r[2], "timestamp": r[3]} for r in cursor.fetchall()]
    conn.close()
    return messages

@app.put("/api/chat/sessions/{session_id}")
async def rename_session(session_id: str, request: RenameRequest, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Verify ownership
    cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (session_id,))
    row = cursor.fetchone()
    if not row or row[0] != current_user['id']:
        conn.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    cursor.execute('UPDATE chat_sessions SET title = ? WHERE id = ?', (request.title, session_id))
    conn.commit()
    conn.close()
    return {"status": "success", "title": request.title}

@app.delete("/api/chat/sessions/{session_id}")
async def delete_session(session_id: str, current_user: dict = Depends(get_current_user)):
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Verify ownership
    cursor.execute('SELECT user_id FROM chat_sessions WHERE id = ?', (session_id,))
    row = cursor.fetchone()
    if not row or row[0] != current_user['id']:
        conn.close()
        raise HTTPException(status_code=404, detail="Session not found")
        
    # Cascade delete messages (if foreign key support enabled, otherwise manual)
    cursor.execute('DELETE FROM chat_messages WHERE session_id = ?', (session_id,))
    cursor.execute('DELETE FROM chat_sessions WHERE id = ?', (session_id,))
    conn.commit()
    conn.close()
    return {"status": "success"}

# Helper function to get conversation history
def get_conversation_history(session_id: str, limit: int = 5) -> List[dict]:
    """Get last N messages from a session for context"""
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT role, content FROM chat_messages 
        WHERE session_id = ? 
        ORDER BY timestamp DESC 
        LIMIT ?
    ''', (session_id, limit))
    messages = [{'role': r[0], 'content': r[1]} for r in cursor.fetchall()]
    conn.close()
    return list(reversed(messages))  # Reverse to get chronological order


# --- Memory & Context Management ---

# Global Upload State Tracking
upload_states: Dict[str, Dict[str, Any]] = {}

def index_upload_background(file_id: str, filename: str, content: str):
    global upload_states
    try:
        upload_states[file_id] = {"status": "indexing", "progress": 0, "filename": filename}
        
        embed_model_name = "nomic-embed-text-v1.5.f16.gguf"
        embed_model_path = MODELS_DIR / embed_model_name
        
        if not embed_model_path.exists():
            logger.error("Embedding model not found for upload indexing")
            upload_states[file_id] = {"status": "error", "error": "Model missing"}
            return
            
        # Use thread lock for model loading safety
        manager = model_manager
        
        # Check if model load is needed first to update status?
        # Actually initializing embedding_fn might take a moment if not loaded
        
        embedding_fn = GGUFEmbeddingFunction(str(embed_model_path))
        client = get_chroma_client()
        collection = client.get_or_create_collection(name="temp_uploads", embedding_function=embedding_fn)
        
        # Chunking
        upload_states[file_id]["status"] = "chunking"
        # Reduce chunk size for safer processing during upload
        chunks = recursive_character_text_splitter(content, chunk_size=300, chunk_overlap=50)
        
        # Limit chunks to avoid overly long processing for large files on uploading
        if len(chunks) > 100:
             chunks = chunks[:100]

        total_chunks = len(chunks)
        if total_chunks == 0:
            upload_states[file_id] = {"status": "ready", "progress": 100}
            return

        ids = []
        metadatas = []
        docs = []
        
        upload_states[file_id]["status"] = "embedding"
        
        # Process in smaller batches to update progress
        batch_size = 10
        for i in range(0, total_chunks, batch_size):
            batch_end = min(i + batch_size, total_chunks)
            batch_chunks = chunks[i:batch_end]
            
            # Prepare data
            for j, chunk in enumerate(batch_chunks):
                abs_index = i + j
                ids.append(f"{file_id}_{abs_index}")
                metadatas.append({"file_id": file_id, "filename": filename, "chunk_index": abs_index})
                docs.append(chunk)
            
            # Add to chroma (this triggers embedding which is slow)
            try:
                # We need to make sure we don't hold the lock for too long if we used one, 
                # but Chromadb/GGUFEmbeddingFn handles the model access via our manager which has a lock.
                collection.add(ids=ids[i:batch_end], documents=batch_chunks, metadatas=metadatas[i:batch_end])
            except Exception as e:
                logger.error(f"Error adding batch {i}: {e}")
            
            # Update progress
            progress = int((batch_end / total_chunks) * 100)
            upload_states[file_id]["progress"] = progress
            
        logger.info(f"Successfully indexed {len(chunks)} chunks for {filename}")
        upload_states[file_id] = {"status": "ready", "progress": 100}
            
    except Exception as e:
        logger.error(f"Background indexing upload error: {e}")
        upload_states[file_id] = {"status": "error", "error": str(e)}

@app.get("/api/chat/file/{file_id}/status")
async def get_upload_status(file_id: str, current_user: dict = Depends(get_current_user)):
    status = upload_states.get(file_id)
    if not status:
        return {"status": "unknown"}
    return status

@app.post("/api/chat/upload")
async def upload_file_context(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...), 
    user: dict = Depends(get_current_user)
):
    try:
        content = ""
        filename = file.filename
        import uuid
        file_id = str(uuid.uuid4())
        
        # Save temp file
        temp_path = Path(f"/tmp/{file_id}_{filename}")
        with open(temp_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        file_path = temp_path
        
        # Extract text based on extension
        if filename.endswith('.docx'):
            content = read_docx_file(file_path)
        elif filename.endswith('.xlsx'):
            content = read_excel_file(file_path)
        elif filename.endswith('.pdf'):
            content = read_pdf_file(file_path)
        else:
            # Try reading as text with fallback encodings
            encodings = ['utf-8', 'shift_jis', 'latin-1']
            for enc in encodings:
                try:
                    content = file_path.read_text(encoding=enc)
                    if content:
                        break
                except:
                    continue
        
        if file_path.exists():
            file_path.unlink()
            
        logger.info(f"Uploaded file {filename}, extracted content length: {len(content)}")
            
        if not content:
            logger.warning(f"Empty content for file: {filename}")
            return JSONResponse(status_code=400, content={"error": "Could not extract text or file is empty"})
            
        # Initialize state
        upload_states[file_id] = {"status": "queued", "progress": 0}

        # Index Logic for Uploaded File (Background)
        background_tasks.add_task(index_upload_background, file_id, filename, content)

        return {"file_id": file_id, "filename": filename}
        
    except Exception as e:
        logger.error(f"Upload error: {e}")
        return {"error": str(e)}

# --- Memory & Context Management ---

def get_recent_summaries(session_id: str, limit: int = 5) -> str:
    """Retrieve recent summaries for the current conversation flow"""
    # Logic: Get summaries for THIS session first (older parts of current conv)
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        SELECT summary FROM chat_summaries 
        WHERE session_id = ? 
        ORDER BY created_at ASC 
    ''', (session_id,))
    rows = cursor.fetchall()
    conn.close()
    
    if not rows:
        return ""
        
    summary_block = "Conversation Summaries (Previous context of this conversation):\n"
    for r in rows:
        summary_block += f"- {r[0]}\n"
    return summary_block + "\n"

async def summarize_old_messages(session_id: str, model_path: Path):
    """
    Background Task: Check if message count > threshold, then summarize oldest chunk and move to summaries table.
    """
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    cursor = conn.cursor()
    
    # Check total message count
    cursor.execute("SELECT count(*) FROM chat_messages WHERE session_id = ?", (session_id,))
    count = cursor.fetchone()[0]
    
    # Thresholds
    MAX_MESSAGES = 20 # Keep last ~20 messages (approx 10 turns) in raw log
    CHUNK_SIZE = 10 # Number of messages to summarize at once
    
    if count <= MAX_MESSAGES:
        conn.close()
        return

    # Find last summarized point (by checking max range_end in summaries)
    cursor.execute("SELECT MAX(range_end) FROM chat_summaries WHERE session_id = ?", (session_id,))
    last_end = cursor.fetchone()[0]
    if last_end is None:
        last_end = 0
        
    # Get next batch of messages to summarize
    # We exclude the most recent MAX_MESSAGES from summarization to keep them as "raw context"
    
    # Let's count how many unwrapped messages exist
    cursor.execute("SELECT id, role, content FROM chat_messages WHERE session_id = ? AND id > ? ORDER BY id ASC", (session_id, last_end))
    all_unsummed = cursor.fetchall()
    
    # If the number of unsummarized messages > MAX_MESSAGES, we summarize the overflow
    users_overflow_count = len(all_unsummed) - MAX_MESSAGES
    
    if users_overflow_count < CHUNK_SIZE:
        conn.close()
        return # Not enough overflow to make a chunk
        
    # Take the first CHUNK_SIZE messages from the unsummed list
    chunk_to_process = all_unsummed[:CHUNK_SIZE]
    
    range_start = chunk_to_process[0][0]
    range_end = chunk_to_process[-1][0]
    
    text_to_summary = "\n".join([f"{r[1]}: {r[2]}" for r in chunk_to_process])
    
    conn.close() # Close mainly to release db for a moment
    
    # Generate Summary using LLM
    try:
        # Use Fast model for summarization
        fast_model = MODELS_DIR / "qwen2-1.5b-instruct-q8_0.gguf"
        if not fast_model.exists():
            if model_path.exists():
                fast_model = model_path
            else:
                return 

        llm = model_manager.get_llm(str(fast_model))
        
        prompt = f"""Summarize the following conversation segment concisely in 2-3 sentences. Capture key facts and topics.

{text_to_summary}

Summary:"""

        output = llm.create_completion(
            prompt=prompt,
            max_tokens=200,
            stop=["\n\n"]
        )
        summary = output['choices'][0]['text'].strip()
        
        # Save to DB
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO chat_summaries (session_id, summary, range_start, range_end)
            VALUES (?, ?, ?, ?)
        ''', (session_id, summary, range_start, range_end))
        conn.commit()
        conn.close()
        
        logger.info(f"Generated summary for session {session_id} (msgs {range_start}-{range_end})")
        
    except Exception as e:
        logger.error(f"Summarization failed: {e}")

# Streaming Chat Endpoint with Memory Architecture
@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest, background_tasks: BackgroundTasks, current_user: dict = Depends(get_current_user)):
    # Disable RAG if indexing (same as before)
    use_nas_override = request.use_nas
    system_notice = ""
    if state.is_indexing and request.use_nas:
        use_nas_override = False
        system_notice = "\n[System Note: Indexing is in progress. Database search is temporarily disabled.]"
    
    async def generate():
        try:
            # 1. Model Selection & Context Setup
            user_models_dir = ensure_user_models_dir(current_user['username'])
            
            model_filename = request.model_id
            model_filename = request.model_id
            if request.model_id == "Fast" or request.model_id == "2":
                model_filename = "qwen2.5-3b-instruct-q4_0.gguf"
            elif request.model_id == "Thinking" or request.model_id == "1":
                 # Removed "Thinking" model as per request, fallback to default
                 model_filename = "qwen2.5-3b-instruct-q4_0.gguf"

            
            model_path = user_models_dir / model_filename
            if not model_path.exists():
                 # Fallback logic
                if (user_models_dir / request.model_id).exists():
                    model_path = user_models_dir / request.model_id
                else:
                    ggufs = [p for p in user_models_dir.glob("*.gguf") if "embed" not in p.name.lower()]
                    if ggufs:
                        model_path = ggufs[0]
                    else:
                        yield 'data: {\"error\": \"No chat models found\"}\\n\\n'
                        return

            # --- AGENT ROUTING ---
            if request.model_id.lower() == "agent" or request.model_id == "秘書モード":
                if not agent_gateway:
                    yield f"data: {json.dumps({'error': 'Agent system not initialized'})}\n\n"
                    return

                # Auto-initialize Agent if needed (Lazy Load)
                if not agent_gateway.agent:
                     conn_agent = sqlite3.connect(DB_PATH)
                     cursor_agent = conn_agent.cursor()
                     cursor_agent.execute("SELECT value FROM settings WHERE key = 'agent_reflex_model'")
                     r = cursor_agent.fetchone()
                     cursor_agent.execute("SELECT value FROM settings WHERE key = 'agent_planner_model'")
                     p = cursor_agent.fetchone()
                     conn_agent.close()
                     
                     # Force default to qwen2.5-3b if not configured or files missing
                     default_model = MODELS_DIR / "qwen2.5-3b-instruct-q4_0.gguf"
                     
                     r_path_str = r[0] if r else str(default_model)
                     p_path_str = p[0] if p else str(default_model)
                     
                     r_path = MODELS_DIR / r_path_str
                     p_path = MODELS_DIR / p_path_str
                     
                     # Construct absolute paths if they are just filenames
                     if not r_path.is_absolute(): r_path = MODELS_DIR / r_path.name
                     if not p_path.is_absolute(): p_path = MODELS_DIR / p_path.name

                     if not r_path.exists(): r_path = default_model
                     if not p_path.exists(): p_path = default_model

                     # Define DB Handlers
                     def db_read_canvas(canvas_id: str):
                         try:
                             conn = sqlite3.connect(DB_PATH)
                             cur = conn.cursor()
                             cur.execute("SELECT content, language, title FROM canvases WHERE id = ?", (canvas_id,))
                             row = cur.fetchone()
                             conn.close()
                             if row:
                                 lang = row[1]
                                 return {"content": row[0], "language": lang, "title": row[2], "template": "document" if lang == "document" else "none"} 
                             return None
                         except Exception as e:
                             print(f"DB Read Error: {e}")
                             return None
                     
                     def db_save_canvas(canvas_id: str, session_id: str, title: str, content: str, language: str):
                         try:
                             conn = sqlite3.connect(DB_PATH)
                             cur = conn.cursor()
                             
                             # App Separation Logic: Docs stored in persistent storage specific to the user
                             target_session_id = session_id
                             # Check if it's a doc (language or content heuristic could apply, but 'document' is explicit)
                             if language == 'document':
                                 # Find owner of current session
                                 cur.execute("SELECT user_id FROM chat_sessions WHERE id = ?", (session_id,))
                                 row = cur.fetchone()
                                 if row:
                                     user_id = row[0]
                                     # Use a persistent session for this user's docs
                                     target_session_id = f"docs_storage_{user_id}"
                                     # Ensure persistent session exists
                                     cur.execute("INSERT OR IGNORE INTO chat_sessions (id, user_id, title) VALUES (?, ?, ?)", 
                                                 (target_session_id, user_id, "My Docs"))
                             
                             now = datetime.utcnow().isoformat()
                             cur.execute("""
                                 INSERT INTO canvases (id, session_id, title, content, language, created_at, updated_at)
                                 VALUES (?, ?, ?, ?, ?, ?, ?)
                                 ON CONFLICT(id) DO UPDATE SET
                                     session_id=excluded.session_id,
                                     title=excluded.title,
                                     content=excluded.content,
                                     language=excluded.language,
                                     updated_at=excluded.updated_at
                             """, (canvas_id, target_session_id, title, content, language, now, now))
                             conn.commit()
                             conn.close()
                         except Exception as e:
                             print(f"DB Save Error: {e}")

                     def db_save_memory(key: str, value: str):
                         try:
                             conn = sqlite3.connect(DB_PATH)
                             cur = conn.cursor()
                             cur.execute("INSERT OR REPLACE INTO user_memory (user_id, key, value) VALUES (?, ?, ?)", 
                                        (current_user["id"], key, value))
                             conn.commit()
                             conn.close()
                         except Exception as e:
                             print(f"Memory Save Error: {e}")

                     if r_path.exists():
                         agent_gateway.initialize(str(r_path), str(p_path), 
                                               db_handler=db_read_canvas, 
                                               db_save_handler=db_save_canvas,
                                               db_memory_handler=db_save_memory)
                     else:
                         yield f"data: {json.dumps({'error': 'Agent models missing (qwen2.5-3b). Check models directory.'})}\n\n"
                         return

                session_id = request.session_id
                if not session_id:
                     import uuid
                     session_id = str(uuid.uuid4())
                     
                     # Create session record
                     conn = sqlite3.connect(DB_PATH, check_same_thread=False)
                     cursor = conn.cursor()
                     cursor.execute('INSERT INTO chat_sessions (id, user_id, title) VALUES (?, ?, ?)', 
                                   (session_id, current_user['id'], request.message[:20]))
                     conn.commit()
                     conn.close()

                # Save User Message
                conn = sqlite3.connect(DB_PATH, check_same_thread=False)
                cursor = conn.cursor()
                ts = datetime.utcnow().isoformat()
                cursor.execute('INSERT INTO chat_messages (session_id, role, content, timestamp) VALUES (?, ?, ?, ?)', 
                              (session_id, 'user', request.message, ts))
                conn.commit()
                conn.close()

                # --- Run Agent Loop with Streaming ---
                agent = agent_gateway.agent
                if not session_id in agent_gateway.sessions:
                    from agent_core import AgentContext
                    agent_gateway.sessions[session_id] = AgentContext(session_id=session_id)
                
                context = agent_gateway.sessions[session_id]
                
                # Execute the new Async Generator Loop
                try:
                    async for event in agent.run_solo_loop(context, request.message):
                        if "status" in event:
                             yield f"data: {json.dumps({'status': event['status']})}\n\n"
                        
                        if "thought_chunk" in event:
                             # Direct streaming to user's view
                             # We use a custom content wrapper or just 'content' 
                             # The frontend appends 'content' to the last message.
                             # To distinguish "Thought" from "Final Answer", we might want a prefix?
                             # But ReAct output is mixed. Let's just stream it.
                             # If we want a separate "Thought Box", frontend needs support.
                             # For now, stream as content, but maybe blockquoted?
                             # Real-time blockquoting is hard.
                             # Let's stream as raw content. The user will see:
                             # "Thought: ... Action: ..."
                             chunk = event['thought_chunk']
                             yield f"data: {json.dumps({'content': chunk})}\n\n"

                        if "thought" in event:
                             # Deprecated, but keep for fallback
                             pass
                        
                        if "action" in event:
                             action_md = f"\n> 🛠️ **Action**: `{event['action']}`\n> Args: `{event['input']}`\n\n"
                             status_msg = f"Executing {event['action']}..."
                             yield f"data: {json.dumps({'content': action_md, 'status': status_msg})}\n\n"

                        if "canvas_update" in event:
                             c_data = event['canvas_update']
                             yield f"data: {json.dumps({'canvas_content': c_data.get('content', ''), 'canvas_language': c_data.get('language', 'html'), 'canvas_title': c_data.get('title', '')})}\n\n"
                        
                        if "observation" in event:
                             result = event['observation']
                             # Check for Canvas Update
                             canvas_match = re.search(r'\[CANVAS_UPDATE\](.*?)\[/CANVAS_UPDATE\]', result, re.DOTALL)
                             if canvas_match:
                                 try:
                                     canvas_data = json.loads(canvas_match.group(1))
                                     yield f"data: {json.dumps({'canvas_content': canvas_data.get('content', ''), 'canvas_language': canvas_data.get('language', 'html')})}\n\n"
                                     # Remove marker
                                     result = re.sub(r'\[CANVAS_UPDATE\].*?\[/CANVAS_UPDATE\]', '', result, flags=re.DOTALL).strip()
                                 except: pass
                             
                             obs_md = f"\n> 🔍 **Observation**\n> {result[:800]}\n\n"
                             yield f"data: {json.dumps({'content': obs_md, 'status': 'Analyzed.'})}\n\n"

                        if "final" in event:
                             final_ans = event['final']
                             # Do NOT yield content again, as it was streamed incrementally.
                             # yield f"data: {json.dumps({'content': final_ans})}\n\n"
                             
                             # Save Assistant Message
                             conn = sqlite3.connect(DB_PATH, check_same_thread=False)
                             cursor = conn.cursor()
                             ts = datetime.utcnow().isoformat()
                             cursor.execute('INSERT INTO chat_messages (session_id, role, content, timestamp) VALUES (?, ?, ?, ?)', 
                                           (session_id, 'assistant', final_ans, ts))
                             conn.commit()
                             conn.close()
                             
                             yield f"data: {json.dumps({'session_id': session_id, 'done': True})}\n\n"
                             return

                except Exception as e:
                    logger.error(f"Agent Loop Error: {e}")
                    error_json = json.dumps({'content': f'\n\n[System Error: {str(e)}]\n\n'})
                    yield f"data: {error_json}\n\n"
                    return

            session_id = request.session_id
            new_session = False
            user_id = current_user['id']
            
            # Create session if needed (to get ID for memory retrieval)
            conn = sqlite3.connect(DB_PATH, check_same_thread=False)
            cursor = conn.cursor()
            
            if not session_id:
                import uuid
                session_id = str(uuid.uuid4())
                title = request.message[:20]
                cursor.execute('INSERT INTO chat_sessions (id, user_id, title) VALUES (?, ?, ?)', 
                              (session_id, user_id, title))
                new_session = True
                conn.commit()
            
            # Save User Message immediately
            ts = datetime.utcnow().isoformat()
            cursor.execute('INSERT INTO chat_messages (session_id, role, content, timestamp) VALUES (?, ?, ?, ?)', 
                          (session_id, 'user', request.message, ts))
            conn.commit()
            
            conn.close()

            # 2. Build Context (The 3 Layers now: System, Summaries, Log)
            
            # Layer 1: System Prompt (Base)
            if use_nas_override or request.attached_file_ids:
                base_system_prompt = (
                    "あなたは優秀なAIアシスタントです。ユーザーの質問に対し、提供された【参照資料】（添付ファイルやデータベース検索結果）の内容を基に回答してください。\n"
                    "・資料の中に答えがある場合は、その部分を引用・要約して答えてください。\n"
                    "・資料に関連情報がない場合は、その旨を伝えてください。\n"
                    "・ユーザーの質問が特定の抽出タスク（例：「会社名を挙げて」など）であればそれに従いますが、基本的には資料の内容についての解説や要約、質問への回答を行ってください。\n"
                    "・重要: ユーザーが使用している言語（日本語、英語など）で回答してください。"
                )
            else:
                base_system_prompt = "あなたは親切なAIアシスタントです。ユーザーの質問と同じ言語で丁寧に答えてください。日本語の質問には日本語で、英語の質問には英語で回答してください。"

            if request.canvas_mode:
                base_system_prompt += (
                    "\n\n【Canvasモード】\n"
                    "あなたはユーザーの依頼に基づいてコードや文書を作成できるAIアシスタントです。\n"
                    "通常の会話も自由に行えますが、コードや文書を生成する際は以下のルールを守ってください：\n"
                    "\n"
                    "1. コードや文書を作成する場合は、必ず<<<CANVAS_START>>>タグを使用すること\n"
                    "2. 通常のマークダウンコードブロック（```）は使用せず、Canvasタグを使うこと\n"
                    "3. Canvasタグの前後で説明や会話をしても構いません\n"
                    "\n"
                    "出力例:\n"
                    "User: Pythonで五目並べを作って\n"
                    "Assistant: 五目並べのゲームを作成しますね。\n"
                    "<<<CANVAS_START>>>\n"
                    "Title: 五目並べゲーム\n"
                    "Language: python\n"
                    "<<<CONTENT_START>>>\n"
                    "# ゲームコード\n"
                    "print('五目並べ')\n"
                    "<<<CANVAS_END>>>\n"
                    "このコードで基本的な五目並べが動作します。"
                )
            
            # Layer 2: Conversation Summaries (Long-term context of this thread)
            summaries_block = get_recent_summaries(session_id)
            
            # Layer 3: Current Conversation Log (Sliding Window)
            current_log_messages = get_conversation_history(session_id, limit=15) 
            
            # Layer 1.5: RAG Context (NAS) & Attached Files
            nas_context = ""
            
            # Attached Files RAG
            if request.attached_file_ids:
                yield f"data: {json.dumps({'status': '添付ファイルを分析中...'})}\n\n"
                await asyncio.sleep(0)
                try:
                    client = get_chroma_client()
                    collection = client.get_collection("temp_uploads") # Assuming it exists if IDs are passed
                    
                    embed_model_name = "nomic-embed-text-v1.5.f16.gguf"
                    embed_model_path = user_models_dir / embed_model_name
                    
                    if embed_model_path.exists():
                        # Run RAG in thread pool to avoid blocking async loop
                        from starlette.concurrency import run_in_threadpool
                        
                        def perform_rag():
                            # Retrieve content directly instead of semantic search
                            # This ensures we get the actual file content regardless of the query
                            return collection.get(
                                where={"file_id": {"$in": request.attached_file_ids}},
                                limit=5, # Reduced from 15 to 5 to fit in 4096 context safely
                                include=['documents', 'metadatas']
                            )

                        results = await run_in_threadpool(perform_rag)
                        
                        if results['documents']:
                            # collection.get returns flat lists, unlike collection.query
                            doc_texts = results['documents']
                            metas = results['metadatas']
                            nas_context += "--- 添付ファイル分析対象 ---\n"
                            for i, text in enumerate(doc_texts):
                                m = metas[i]
                                nas_context += f"【添付データ NO.{i+1}】\n"
                                nas_context += "[[本文開始]]\n"
                                nas_context += f"{text}\n"
                                nas_context += "[[本文終了]]\n"
                                nas_context += f"（※ファイル名: {m.get('filename', 'Unknown')}）\n"
                                nas_context += "---------------------------------\n"
                            nas_context += "--- 添付ファイル終了 ---\n\n"
                except Exception as e:
                     logger.error(f"Attached File RAG Error: {e}")

            if use_nas_override:
                yield f"data: {json.dumps({'status': 'データベースを検索中...'})}\n\n"
                await asyncio.sleep(0)
                try:
                    storage_mode = get_storage_mode()
                    collection_name = f"documents_{storage_mode}"
                    client = get_chroma_client()
                    try:
                        collection = client.get_collection(collection_name)
                    except:
                        collection = None

                    if collection:
                        embed_model_name = "nomic-embed-text-v1.5.f16.gguf"
                        embed_model_path = user_models_dir / embed_model_name
                        if embed_model_path.exists():
                            embedding_fn = GGUFEmbeddingFunction(str(embed_model_path))
                            # Add prefix for better retrieval with nomic
                            # Enhanced intent-based query construction
                            # Extract nouns/keywords from message for vector search
                            keywords = " ".join(re.findall(r'[一-龠ぁ-んァ-ヶa-zA-Z0-9]+', request.message))
                            query_text = f"search_query: {keywords}"
                            query_embed = embedding_fn([query_text])[0]
                            
                            # Yield heartbeat before long search
                            yield f"data: {json.dumps({'status': '最良の資料を抽出中...'})}\n\n"
                            await asyncio.sleep(0.01)

                            # Safe number of results for 8k context window
                            results = collection.query(query_embeddings=[query_embed], n_results=12)
                            
                            if results['documents']:
                                doc_texts = results['documents'][0]
                                metas = results['metadatas'][0]
                                logger.info(f"RAG: Found {len(doc_texts)} relevant chunks from {collection_name}")
                                
                                nas_context += "\n--- 分析対象データ・セット開始 ---\n"
                                for i, text in enumerate(doc_texts):
                                    m = metas[i]
                                    nas_context += f"【データ NO.{i+1}】\n"
                                    nas_context += "[[本文開始]]\n"
                                    nas_context += f"{text}\n"
                                    nas_context += "[[本文終了]]\n"
                                    nas_context += f"（※このデータの出典ファイル: {m.get('filename', 'Unknown')}）\n"
                                    nas_context += "---------------------------------\n"
                                nas_context += "--- 分析対象データ・セット終了 ---\n\n"
                            else:
                                logger.info(f"RAG: No relevant chunks found in {collection_name}")
                except Exception as e:
                    logger.error(f"RAG Error: {e}")

            # Assemble Final Messages List
            yield f"data: {json.dumps({'status': '思考を整理中...'})}\n\n"
            await asyncio.sleep(0.05)
            final_messages = []
            
            # System Message Construction
            full_system_content = base_system_prompt + "\n\n"
            if summaries_block:
                full_system_content += summaries_block + "\n"
            if nas_context:
                # Safety Truncate nas_context if it looks too big
                # With n_ctx=4096, assume approx 3000 chars for context + 1000 for log/system is safe-ish
                if len(nas_context) > 3500:
                    logger.warning("Truncating NAS context to prevent overflow")
                    nas_context = nas_context[:3500] + "\n...(truncated)..."

                full_system_content += "=== 参照資料 ===\n" + nas_context + "\n"
            if system_notice:
                full_system_content += system_notice

            final_messages.append({"role": "system", "content": full_system_content})
            
            # Append Chat Log
            final_messages.extend(current_log_messages)
            
            # FORCE REMINDER for Canvas Mode
            if request.canvas_mode:
                final_messages.append({
                    "role": "system", 
                    "content": "REMINDER: You are in Canvas Mode. DO NOT use markdown code blocks. OUTPUT ONLY using <<<CANVAS_START>>> tags."
                })
            
            yield f"data: {json.dumps({'status': '考察中...'})}\n\n"
            await asyncio.sleep(0.05)
            
            # logger.debug(f"Prompt: {final_messages}") # Too noisy
            
            # Append Current User Message
            # Append Current User Message (Already in log if saved to DB, but get_conversation_history logic might exclude it if limit reached or if explicitly coded. 
            # get_conversation_history gets last N. If we just inserted it, it IS the last one.
            # So current_log_messages INCLUDES it.
            # We should NOT append it again.
            # final_messages.append({"role": "user", "content": request.message}) 
            # Removing the line.

            # 3. Generate Response
            # --- Canvas Agent Logic (Simplified) ---
            # Bypass complex Manager/Worker split for now to ensure reliability with smaller models
            if request.canvas_mode:
                 # Check for dedicated code model if available, otherwise use default
                 # For now, just use the selected model but with the strict system prompt applied above
                 llm = model_manager.get_llm(str(model_path))
            else:
                llm = model_manager.get_llm(str(model_path))
            
            # Yield a small pulse to keep connection alive during prefill
            yield f"data: {json.dumps({'content': '', 'status': '考え中...'})}\n\n"
            await asyncio.sleep(0.1)

            full_response = ""
            logger.info(f"Starting LLM generation. Prompt chars: {len(str(final_messages))}")
            
            # Stream Interceptor State
            buffer = ""
            inside_canvas = False
            
            try:
                for chunk in llm.create_chat_completion(
                    messages=final_messages,
                    max_tokens=2048,
                    temperature=0.7,
                    stream=True
                ):
                    if 'choices' in chunk and len(chunk['choices']) > 0:
                        delta = chunk['choices'][0].get('delta', {})
                        if 'content' in delta:
                            content = delta['content']
                            
                            if request.canvas_mode:
                                # Buffer logic to detect backticks across chunks
                                buffer += content
                                
                                # Process buffer for start/end tags
                                # We look for the triple backtick marker
                                while '```' in buffer:
                                    # Split at the first occurrence
                                    pre, sep, post = buffer.partition('```')
                                    
                                    if not inside_canvas:
                                        # Opening code block
                                        # Check if 'post' has the language identifier (e.g. "python\n")
                                        # We need to peek a bit ahead to get the language
                                        if '\n' in post or len(post) > 15: 
                                            # Extract language
                                            lang_match = re.match(r'^([a-zA-Z0-9+#]*)\s*', post)
                                            lang = lang_match.group(1) if lang_match else ""
                                            
                                            # Remove language from 'post'
                                            consumed_len = len(lang_match.group(0)) if lang_match else 0
                                            post = post[consumed_len:]
                                            
                                            # Generate Canvas Header
                                            # Note: We strip the ``` and language, replacing with Tag
                                            replacement = (
                                                f"\n<<<CANVAS_START>>>\n"
                                                f"Title: Generated Code\n"
                                                f"Language: {lang or 'text'}\n"
                                                f"<<<CONTENT_START>>>\n"
                                            )
                                            
                                            yield f"data: {json.dumps({'content': pre + replacement})}\n\n"
                                            full_response += pre + replacement
                                            
                                            inside_canvas = True
                                            buffer = post # Continue processing 'post' in loop
                                        else:
                                            # Not enough data to determine language yet, wait for more chunks
                                            break
                                            
                                    else:
                                        # Closing code block (already inside)
                                        # Just replace ``` with END tag
                                        replacement = "\n<<<CANVAS_END>>>\n"
                                        
                                        yield f"data: {json.dumps({'content': pre + replacement})}\n\n"
                                        full_response += pre + replacement
                                        
                                        inside_canvas = False
                                        buffer = post # Continue processing
                                
                                # If we broke out of loop (waiting for more data), or finished
                                # We DO NOT yield the buffer yet if we are waiting for lang check
                                # But we must ensure we don't hold text forever.
                                # Simple heuristic: if buffer doesn't contain partial backticks, yield it.
                                if '`' not in buffer:
                                     yield f"data: {json.dumps({'content': buffer})}\n\n"
                                     full_response += buffer
                                     buffer = ""
                                else:
                                    # Keep safe amount, yield safe header
                                    # If buffer grows too big without completing backtick, flush it
                                    if len(buffer) > 50:
                                        # Flush safe part
                                        cutoff = buffer.rfind('`')
                                        if cutoff == -1: cutoff = len(buffer)
                                        safe_part = buffer[:cutoff]
                                        keep_part = buffer[cutoff:]
                                        
                                        yield f"data: {json.dumps({'content': safe_part})}\n\n"
                                        full_response += safe_part
                                        buffer = keep_part
                            else:
                                # Normal Mode
                                full_response += content
                                yield f"data: {json.dumps({'content': content})}\n\n"
                                await asyncio.sleep(0.01)

                # Flush remaining buffer at end of stream
                if request.canvas_mode and buffer:
                    yield f"data: {json.dumps({'content': buffer})}\n\n"
                    full_response += buffer

            except Exception as e:
                logger.error(f"Streaming Error: {e}")
                yield f"data: {json.dumps({'error': str(e)})}\n\n"
            
            logger.info(f"Generation complete. Response length: {len(full_response)}")
            
            # 4. Save & Post-Processing
            conn = sqlite3.connect(DB_PATH, check_same_thread=False)
            cursor = conn.cursor()
            
            # Update session timestamp
            cursor.execute('UPDATE chat_sessions SET updated_at = ? WHERE id = ?', 
                          (datetime.utcnow().isoformat(), session_id))
                          
            # Save messages
            ts = datetime.utcnow().isoformat()
            # User message already saved
            cursor.execute('INSERT INTO chat_messages (session_id, role, content, timestamp) VALUES (?, ?, ?, ?)', 
                          (session_id, 'assistant', full_response, ts))
                          
            conn.commit()
            conn.close()
            
            # Trigger Background Summarization
            background_tasks.add_task(summarize_old_messages, session_id, model_path)
            
            yield f"data: {json.dumps({'session_id': session_id, 'title': request.message[:20] if new_session else None, 'done': True})}\n\n"

        except Exception as e:
            logger.error(f"Streaming Error: {e}")
            yield 'data: {\"error\": \"' + str(e) + '\"}\\n\\n'

    return StreamingResponse(generate(), media_type="text/event-stream")

# Legacy Chat Endpoint (Redirect to use logic if needed, but for now we just keep it simple or deprecate)
@app.post("/api/chat")
async def chat(request: ChatRequest):
    return {"error": "Please use /api/chat/stream endpoint for this version"}

# --- NAS File Explorer Endpoints ---

class NASFile(BaseModel):
    name: str
    path: str
    is_dir: bool
    size: Optional[int] = None
    last_modified: Optional[str] = None

@app.get("/api/nas/list", response_model=List[NASFile])
async def list_nas_files(path: str = "", source: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    try:
        # Determine root based on source param or global storage mode
        if source == "nas":
            root_dir = MNT_DIR
        elif source == "internal":
            root_dir = INTERNAL_NAS_DIR
        else:
            storage_mode = get_storage_mode()
            root_dir = MNT_DIR if storage_mode == "nas" else INTERNAL_NAS_DIR
        
        # Safe path joining
        target_path = (root_dir / path).resolve()
        if not str(target_path).startswith(str(root_dir.resolve())):
             raise HTTPException(status_code=403, detail="Access denied")
             
        if not target_path.exists():
            return [] # Or raise 404
            
        if not target_path.is_dir():
            raise HTTPException(status_code=400, detail="Not a directory")
            
        files = []
        for item in target_path.iterdir():
            try:
                stat = item.stat()
                # Create relative path for frontend
                rel_path = item.relative_to(root_dir)
                files.append({
                    "name": item.name,
                    "path": str(rel_path),
                    "is_dir": item.is_dir(),
                    "size": stat.st_size if not item.is_dir() else None,
                    "last_modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })
            except Exception as e:
                logger.warning(f"Error accessing {item}: {e}")
                
        # Sort: Directories first, then files
        files.sort(key=lambda x: (not x['is_dir'], x['name'].lower()))
        return files
        
    except Exception as e:
        logger.error(f"List files error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/nas/read")
async def read_nas_file(path: str, source: Optional[str] = None, current_user: dict = Depends(get_current_user)):
    try:
        if source == "nas":
            root_dir = MNT_DIR
        elif source == "internal":
            root_dir = INTERNAL_NAS_DIR
        else:
            storage_mode = get_storage_mode()
            root_dir = MNT_DIR if storage_mode == "nas" else INTERNAL_NAS_DIR

        
        target_path = (root_dir / path).resolve()
        if not str(target_path).startswith(str(root_dir.resolve())):
             raise HTTPException(status_code=403, detail="Access denied")
             
        if not target_path.exists() or not target_path.is_file():
            raise HTTPException(status_code=404, detail="File not found")
            
        # Read content
        content = ""
        if target_path.suffix == '.docx':
            content = read_docx_file(target_path)
        elif target_path.suffix == '.xlsx':
            content = read_excel_file(target_path)
        else:
            # Limit size for safety? 
            if target_path.stat().st_size > 10 * 1024 * 1024:
                raise HTTPException(status_code=400, detail="File too large to read directly (max 10MB)")
            content = target_path.read_text(encoding='utf-8', errors='ignore')
            
        return {"filename": target_path.name, "content": content}
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Read file error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/nas/content")
async def get_nas_content(path: str, source: str = "nas", current_user: dict = Depends(get_current_user)):
    try:
        if source == "nas":
            root_dir = MNT_DIR
        elif source == "internal":
            root_dir = INTERNAL_NAS_DIR
        else:
            storage_mode = get_storage_mode()
            root_dir = MNT_DIR if storage_mode == "nas" else INTERNAL_NAS_DIR
            
        target_path = (root_dir / path).resolve()
        
        # Security check
        if not str(target_path).startswith(str(root_dir.resolve())):
             raise HTTPException(status_code=403, detail="Access denied")
             
        if not target_path.exists() or not target_path.is_file():
            raise HTTPException(status_code=404, detail="File not found")
            
        return FileResponse(target_path)
        
    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Get content error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/nas/upload")
async def upload_nas_file(
    file: UploadFile = File(...), 
    path: str = Form(""), 
    source: str = Form("internal"),
    current_user: dict = Depends(get_current_user)
):
    if source != "internal":
        raise HTTPException(status_code=403, detail="Only internal storage allows uploads")
        
    try:
        root_dir = INTERNAL_NAS_DIR
        target_dir = (root_dir / path).resolve()
        
        if not str(target_dir).startswith(str(root_dir.resolve())):
             raise HTTPException(status_code=403, detail="Access denied")
             
        if not target_dir.exists():
            target_dir.mkdir(parents=True, exist_ok=True)
            
        file_path = target_dir / file.filename
        
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
            
        return {"status": "success", "filename": file.filename}
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/nas/delete")
async def delete_nas_item(path: str, source: str = "internal", current_user: dict = Depends(get_current_user)):
    if source != "internal":
        raise HTTPException(status_code=403, detail="ReadOnly storage")
        
    try:
        root_dir = INTERNAL_NAS_DIR
        target_path = (root_dir / path).resolve()
        
        if not str(target_path).startswith(str(root_dir.resolve())): # Basic jail check
             # One extra check: user shouldn't delete the root drive dir accidentally if path is empty
             if str(target_path) == str(root_dir.resolve()):
                 raise HTTPException(status_code=403, detail="Cannot delete root")
             raise HTTPException(status_code=403, detail="Access denied")
             
        if not target_path.exists():
            raise HTTPException(status_code=404, detail="Not found")
            
        if target_path.is_dir():
            shutil.rmtree(target_path)
        else:
            target_path.unlink()
            
        return {"status": "success"}
    except Exception as e:
        logger.error(f"Delete error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/nas/rename")
async def rename_nas_item(
    item_path: str = Body(...), 
    new_name: str = Body(...), 
    source: str = Body("internal"),
    is_folder_creation: bool = Body(False),
    current_user: dict = Depends(get_current_user)
):
    if source != "internal":
        raise HTTPException(status_code=403, detail="ReadOnly storage")
        
    try:
        root_dir = INTERNAL_NAS_DIR
        
        if is_folder_creation:
            # item_path here is the PARENT directory
            parent_dir = (root_dir / item_path).resolve()
            if not str(parent_dir).startswith(str(root_dir.resolve())):
                 raise HTTPException(status_code=403, detail="Access denied")
            
            new_folder_path = parent_dir / new_name
            new_folder_path.mkdir(exist_ok=True)
            return {"status": "success", "path": str(new_folder_path.relative_to(root_dir))}
        else:
            # Rename existing
            target_path = (root_dir / item_path).resolve()
            if not str(target_path).startswith(str(root_dir.resolve())): 
                 raise HTTPException(status_code=403, detail="Access denied")
                 
            if not target_path.exists():
                raise HTTPException(status_code=404, detail="Not found")
                
            new_path = target_path.parent / new_name
            target_path.rename(new_path)
            return {"status": "success", "path": str(new_path.relative_to(root_dir))}
            
    except Exception as e:
        logger.error(f"Rename/Create error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# -------------------------------------------------------------------
# Drive API (for Memo App & Internal Drive)
# -------------------------------------------------------------------
DRIVE_DIR = INTERNAL_NAS_DIR / "drive"
DRIVE_DIR.mkdir(parents=True, exist_ok=True)

class DriveFile(BaseModel):
    name: str  # Filename only (not full path)
    data: str = ""

@app.get("/api/drive/list")
async def list_drive_files(current_user: dict = Depends(get_current_user)):
    try:
        files = []
        if not DRIVE_DIR.exists():
            return []
        
        for entry in DRIVE_DIR.iterdir():
            if entry.is_file():
                stat = entry.stat()
                files.append({
                    "name": entry.name,
                    "path": str(entry),
                    "size": stat.st_size,
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat()
                })
        
        # Sort by modified desc
        files.sort(key=lambda x: x["modified"], reverse=True)
        return files
    except Exception as e:
        logger.error(f"Error listing drive files: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/drive/read/{filename}")
async def read_drive_file(filename: str, current_user: dict = Depends(get_current_user)):
    try:
        # Sanitize
        safe_name = os.path.basename(filename)
        file_path = DRIVE_DIR / safe_name
        
        if not file_path.exists():
            raise HTTPException(status_code=404, detail="File not found")
            
        return {"content": file_path.read_text(encoding="utf-8")}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/drive/save")
async def save_drive_file(file: DriveFile, current_user: dict = Depends(get_current_user)):
    try:
        if not file.name:
            raise HTTPException(status_code=400, detail="Filename required")
            
        safe_name = os.path.basename(file.name)
        # Ensure extension
        if not safe_name.endswith('.md') and not safe_name.endswith('.txt'):
            safe_name += '.md'
            
        file_path = DRIVE_DIR / safe_name
        file_path.write_text(file.data, encoding="utf-8")
        
        return {"status": "success", "filename": safe_name}
    except Exception as e:
        logger.error(f"Error saving file: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/api/drive/delete/{filename}")
async def delete_drive_file(filename: str, current_user: dict = Depends(get_current_user)):
    try:
        safe_name = os.path.basename(filename)
        file_path = DRIVE_DIR / safe_name
        
        if file_path.exists():
            file_path.unlink()
            return {"status": "success"}
        else:
            raise HTTPException(status_code=404, detail="File not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

# --- License System (Device Flow) ---
LICENSE_FILE = BASE_DIR / ".license"

# Configuration
# Pointing to Next.js API Routes acting as the Mock Portal
PORTAL_BASE_URL = os.getenv("PORTAL_BASE_URL", "https://oonanji-vault.com") 

class LicenseRequest(BaseModel):
    license_key: str

@app.get("/api/license/status")
async def get_license_status():
    if not LICENSE_FILE.exists():
        return {"active": False, "plan": None, "expires_at": None}
    
    try:
        data = json.loads(LICENSE_FILE.read_text())
        if data.get("expires_at"):
            expires = datetime.fromisoformat(data["expires_at"].replace('Z', '+00:00'))
            if datetime.now(timezone.utc) > expires:
                return {"active": False, "error": "Expired", "plan": data.get("plan")}
        
        return {
            "active": data.get("valid", False) or True,
            "plan": data.get("plan"),
            "expires_at": data.get("expires_at"),
            "license_key": data.get("license_key")
        }
    except Exception:
        return {"active": False}

class LoginRequest(BaseModel):
    email: str
    password: str

@app.post("/api/license/login")
async def login_license(req: LoginRequest):
    """Direct Login to Portal"""
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            # Call Portal Login API (Simulated ROPC flow)
            resp = await client.post(f"{PORTAL_BASE_URL}/api/auth/login", json={
                "email": req.email,
                "password": req.password,
                "client_id": "oonanji-vault-client"
            })

            if resp.status_code != 200:
                try:
                    error_data = resp.json()
                    return JSONResponse(status_code=resp.status_code, content=error_data)
                except ValueError:
                     return JSONResponse(status_code=502, content={"error": "portal_error", "detail": "Invalid response from portal"})

            data = resp.json()
            
            # Save license
            license_data = {
                "active": True,
                "valid": True,
                "key": "PORTAL-LINKED-LICENSE",
                "plan": data.get("plan", "standard"),
                "access_token": data.get("access_token"),
                "refresh_token": data.get("refresh_token"),
                "linked_at": datetime.now().isoformat(),
                "email": req.email
            }
            LICENSE_FILE.write_text(json.dumps(license_data))
            
            print(f"[SUCCESS] License Linked for {req.email}. Plan: {license_data['plan']}")
            return {"valid": True, "plan": license_data['plan']}

    except httpx.RequestError as e:
        print(f"[ERROR] Connection failed to {PORTAL_BASE_URL}: {e}")
        return JSONResponse(status_code=503, content={"error": "connection_error", "detail": "Could not connect to license portal"})
    except Exception as e:
        print(f"[ERROR] Login error: {e}")
        return JSONResponse(status_code=500, content={"error": "internal_error"})

# --- License System V2 (Callback Flow) ---

class CallbackRequest(BaseModel):
    code: str

@app.post("/api/license/callback")
async def callback_license(req: CallbackRequest):
    """
    Handle OAuth callback from Portal.
    1. Receive code.
    2. Save it as access token (Simplified for test).
    3. Check license status with Portal.
    """
    try:
        # In a real OAuth flow, we would exchange code for token here.
        # For this test environment, we assume the code IS the token or can be used as one.
        access_token = req.code
        
        # Verify with Portal immediately
        license_status = await check_portal_license(access_token)
        
        # Save license data
        plan = license_status.get("plan", "unknown").lower()
        # Allow update for free/personal plans as well
        allow_update = True  # Enable updates for all authenticated users
        # allow_update = license_status.get("allow_update", False) or plan in ["starter", "enterprise", "corporate_subscribed", "free", "personal"]
        
        license_data = {
            "active": True, # Authenticated
            "valid": True, # If we got here, it's valid enough to be linked
            "key": "PORTAL-LINKED",
            "access_token": access_token,
            "plan": plan,
            "allow_update": allow_update,
            "linked_at": datetime.now().isoformat(),
            "last_checked": datetime.now().isoformat()
        }
        LICENSE_FILE.write_text(json.dumps(license_data))
        
        return {"status": "success", "license": license_data}

    except HTTPException as he:
        raise he
    except Exception as e:
        logger.error(f"Callback error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

async def check_portal_license(token: str):
    """Call Portal API to check license status"""
    url = f"{PORTAL_BASE_URL}/api/check-license"
    logger.info(f"Checking license at {url} with token={token[:5]}...")
    
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            resp = await client.get(url, headers={"Authorization": f"Bearer {token}"})
            
            if resp.status_code == 200:
                data = resp.json()
                logger.info(f"License check success: {data}")
                return data
            else:
                logger.warning(f"License check failed: {resp.status_code} {resp.text}")
                # Fallback for test if Portal is not actually running or mocking
                # raise HTTPException(status_code=401, detail="License check failed")
                return {"allow_update": False, "plan": "free"}
                
    except Exception as e:
        logger.error(f"Failed to connect to Portal: {e}")
        # raise HTTPException(status_code=503, detail="Portal unreachable")
        return {"allow_update": False, "plan": "error"}

@app.post("/api/license/refresh")
async def refresh_license_status():
    """Manually trigger a check against Portal"""
    if not LICENSE_FILE.exists():
         raise HTTPException(status_code=404, detail="No license linked")
         
    try:
        data = json.loads(LICENSE_FILE.read_text())
        token = data.get("access_token")
        if not token:
             raise HTTPException(status_code=400, detail="Invalid license data")
             
        status = await check_portal_license(token)
        
        # Update local state
        data["valid"] = status.get("allow_update", False)
        data["allow_update"] = status.get("allow_update", False)
        data["plan"] = status.get("plan", data.get("plan"))
        data["last_checked"] = datetime.now().isoformat()
        
        LICENSE_FILE.write_text(json.dumps(data))
        return data
        
    except Exception as e:
        logger.error(f"Refresh error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

# Legacy/Manual verify (kept for backward compat or manual fallback)
@app.post("/api/license/verify")
async def verify_license(req: LicenseRequest):
    key = req.license_key.strip()
    if key == "oonanji-dev" or key.startswith("ONJ-") or key.startswith("OONANJI-"):
        mock_response = {
            "valid": True,
            "signature": "mock_jwt_signature_for_dev",
            "plan": "enterprise",
            "expires_at": "2030-12-31T23:59:59Z",
            "license_key": key,
            "allow_update": True
        }
        LICENSE_FILE.write_text(json.dumps(mock_response))
        return mock_response
    raise HTTPException(status_code=400, detail="Invalid license key")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
# trigger reload 2026年  1月 11日 日曜日 16:02:21 JST
